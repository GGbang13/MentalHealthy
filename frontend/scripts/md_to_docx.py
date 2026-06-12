# -*- coding: utf-8 -*-
"""Convert thesis-draft.md to Word .docx (headings, paragraphs, numbered lists, markdown table)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    print("Missing python-docx. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def set_doc_fonts(doc: Document) -> None:
    """Body and headings: Latin TNR, East Asia 宋体, 12pt body."""
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    # 1.5 line spacing for body
    pf = doc.styles["Normal"].paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = text.replace("`", "")
    return text


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|"):
        return False
    inner = s.strip("|").split("|")
    return all(re.fullmatch(r"-+", cell.strip()) for cell in inner if cell.strip())


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip("\n")
        if not line.strip().startswith("|"):
            break
        if is_table_separator(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            val = strip_md_inline(row[ci]) if ci < len(row) else ""
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)


def md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    set_doc_fonts(doc)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Skip markdown table separator rows when not inside a table block
        if stripped.startswith("|") and is_table_separator(stripped):
            i += 1
            continue

        # Markdown table
        if stripped.startswith("|") and not is_table_separator(stripped):
            rows, ni = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            i = ni
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = strip_md_inline(m.group(2).strip())
            if level == 1:
                doc.add_heading(title, level=0)  # Title style for top-level
            else:
                doc.add_heading(title, level=min(level - 1, 9))
            i += 1
            continue

        # Numbered list "1. xxx"
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(strip_md_inline(stripped), style="List Number")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue

        # Bullet "- xxx"
        if stripped.startswith("- ") or stripped.startswith("* "):
            body = strip_md_inline(stripped[2:].lstrip())
            p = doc.add_paragraph(body, style="List Bullet")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            i += 1
            continue

        p = doc.add_paragraph(strip_md_inline(stripped))
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.5
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "thesis-draft.md"
    out = root / "毕业论文初稿-基于SpringBoot和Vue的心理健康服务平台.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not md.is_file():
        print(f"Not found: {md}", file=sys.stderr)
        sys.exit(1)
    md_to_docx(md, out)


if __name__ == "__main__":
    main()
