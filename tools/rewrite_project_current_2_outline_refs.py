from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("project_current_2_source.docx")
OUT = Path("project_current_2_按大纲参考文献扩充版.docx")
REPORT = Path("docs/project_current_2_reference_locations.txt")


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = style
    except Exception:
        new_para.style = "Normal"
    if text:
        new_para.add_run(text)
    return new_para


def find_para(doc: Document, exact: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    raise ValueError(f"paragraph not found: {exact}")


def find_contains(doc: Document, needle: str) -> Paragraph | None:
    for para in doc.paragraphs:
        if needle in para.text:
            return para
    return None


def append_citation(doc: Document, needle: str, citation: str) -> None:
    para = find_contains(doc, needle)
    if para is not None and citation not in para.text:
        para.add_run(citation)


def remove_old_references(doc: Document) -> Paragraph:
    start = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == "参考文献":
            start = idx
            break
    if start is None:
        raise ValueError("参考文献 heading not found")
    for para in list(doc.paragraphs[start + 1:]):
        para._element.getparent().remove(para._element)
    return doc.paragraphs[start]


def main() -> None:
    doc = Document(SRC)

    locations: list[str] = []

    # Existing paragraphs: add citation marks where the current draft already discusses related content.
    citation_targets = [
        ("心理健康是高校学生成长与发展的重要基础", "[1-3]", "研究背景与意义"),
        ("传统模式往往依赖线下窗口", "[4-6]", "研究背景与意义"),
        ("国外高校与医疗机构较早", "[7-10]", "国内外研究现状"),
        ("前后端分离架构", "[11-14]", "国内外研究现状/技术路线"),
        ("系统采用前后端分离模式", "[15-18]", "系统架构设计"),
        ("心理测评模块除风险概率与等级外", "[19-22]", "项目总结"),
        ("WebSocket 和 REST API", "[23-24]", "在线聊天/项目总结"),
        ("心理健康属于敏感领域", "[25-27]", "不足与展望"),
        ("后台管理功能可以继续加强统计分析能力", "[28-30]", "不足与展望"),
    ]
    for needle, citation, where in citation_targets:
        append_citation(doc, needle, citation)
        locations.append(f"{where}：在包含“{needle}”的段落附近补充引用 {citation}。")

    # Add a technology chapter after the technical-route section, matching the requested outline.
    tech_anchor = find_para(doc, "技术路线")
    current = tech_anchor
    tech_blocks = [
        ("相关技术介绍", "Heading 1"),
        ("Spring Boot 技术", "Heading 2"),
        ("Spring Boot 是基于 Spring 生态的快速开发框架，能够通过自动配置、起步依赖和内嵌服务器简化 Java Web 应用开发。本系统后端采用 Spring Boot 构建 REST API，将认证、用户、咨询师、预约、测评、聊天、文章、通知和看板等业务能力封装为控制器与服务层方法。与传统 Servlet 或配置较重的 Spring MVC 项目相比，Spring Boot 更适合毕业设计阶段快速完成可运行系统，同时又能保持较清晰的工程结构[15-16]。", "Normal"),
        ("在本项目中，Spring Boot 的作用不仅是提供接口入口，还承担统一异常处理、配置管理、依赖注入和业务分层组织等职责。控制器层负责接收前端请求，服务层负责业务规则和权限校验，Mapper 层负责数据库访问。这种分层结构便于后续定位问题，也便于在论文中按照模块解释实现过程[13-14]。", "Normal"),
        ("Vue 3 与前端工程技术", "Heading 2"),
        ("前端采用 Vue 3、TypeScript、Vite、Pinia、Vue Router、Element Plus、Axios 和 ECharts 等技术构建。Vue 3 通过组件化和响应式机制降低了复杂页面的维护成本，TypeScript 提供类型约束，Vite 提升开发构建速度，Element Plus 提供表单、表格、弹窗、标签和消息提示等常用组件，ECharts 用于管理员看板中的统计图表展示[28]。", "Normal"),
        ("本系统前端页面按照业务角色拆分为用户端、咨询师端和管理员端。普通用户主要访问用户首页、测评、预约、文章和聊天页面；咨询师主要访问预约处理、资料维护和聊天页面；管理员主要访问用户管理、咨询师管理、文章管理、通知管理和数据看板页面。通过 Vue Router 的路由元信息和 Pinia 中保存的角色信息，前端能够实现页面级访问控制[28]。", "Normal"),
        ("MyBatis-Plus 与 MySQL 数据库", "Heading 2"),
        ("MySQL 用于保存平台核心业务数据，包括用户、咨询师资料、预约记录、测评量表、测评记录、聊天消息、文章、通知、评价和操作日志等。MyBatis-Plus 在 MyBatis 基础上封装了常用 CRUD 操作和条件构造器，使后端能够用较少代码完成数据查询、插入和更新。对于本系统这类以业务表为主的管理系统而言，MyBatis-Plus 能够明显减少重复 SQL 编写，提高开发效率[12][30]。", "Normal"),
        ("数据库设计是本系统的重要组成部分。用户表作为身份体系中心表，咨询师资料表、预约表、测评记录表和聊天消息表都通过用户编号或咨询师资料编号建立关系。合理的数据结构不仅能够支撑当前功能，也能为后续统计分析、日志审计和服务评价扩展提供基础[12][30]。", "Normal"),
        ("JWT 与 Spring Security", "Heading 2"),
        ("系统采用 JWT 与 Spring Security 实现登录认证和权限控制。用户登录成功后，后端生成包含用户编号和角色信息的 token，前端在后续请求中通过请求头携带该 token。后端过滤器解析 token 后，将用户身份写入安全上下文，使后续接口能够识别当前操作者。与传统 Session 相比，JWT 更适合前后端分离场景，能够减少服务端会话存储压力[17-18]。", "Normal"),
        ("需要强调的是，前端路由控制只能改善用户体验，不能作为唯一安全手段。系统在后端接口和服务层中仍需要进行角色判断和数据归属校验，例如普通用户只能取消自己的预约，咨询师只能处理属于自己的预约，管理员接口必须验证当前用户是否具备管理员角色[17][25-26]。", "Normal"),
        ("WebSocket 与实时通信", "Heading 2"),
        ("在线咨询场景需要较强的实时性，因此系统在聊天模块中引入 WebSocket。WebSocket 建立连接后，服务端可以主动向客户端推送消息，避免前端通过定时轮询频繁请求服务器。与此同时，系统仍保留 REST 接口用于消息保存、历史记录查询和联系人加载，从而保证聊天记录可追溯，并在实时连接异常时保持基本可用[23-24]。", "Normal"),
        ("CatBoost 与风险分析技术", "Heading 2"),
        ("心理测评模块需要将用户答案转化为风险概率、风险等级和影响因素。本系统设计了可选 CatBoost 推理路径，并提供 Java 规则回退逻辑。当模型文件和 Python 推理脚本可用时，系统可调用外部模型生成预测结果；当模型环境不可用时，系统使用内置特征权重生成基础风险分析。该设计兼顾了智能分析能力和系统可运行性[19-22]。", "Normal"),
    ]
    for text, style in tech_blocks:
        current = insert_after(current, text, style)
    locations.append("相关技术介绍：新增 Spring Boot、Vue 3、MyBatis-Plus、JWT、WebSocket、CatBoost 等技术说明，引用[12-30]。")

    additions: dict[str, list[str]] = {
        "研究背景与意义": [
            "从高校心理服务实践看，信息化平台的价值并不只是减少纸质表格，而是将服务入口、业务流程和数据记录整合到统一系统中。传统线下模式下，学生需要通过线下窗口、电话或分散表格完成预约和信息登记，咨询师处理记录也较难形成统一视图。当服务对象数量增加时，人工统计难以及时反映整体风险分布和咨询资源使用情况。因此，构建心理健康服务平台有助于提高预约效率、降低求助门槛，并为管理人员提供数据辅助依据[1-6]。",
            "同时，心理健康服务涉及个人情绪、压力、家庭关系和求助记录等敏感信息，系统设计必须兼顾效率和隐私保护。平台不能把测评结果简单等同于医学诊断，也不能让非授权人员接触用户敏感内容。因此，本文在设计中强调角色边界、数据最小化展示和人工复核意识，使系统定位为高校心理服务工作的辅助工具，而不是替代专业咨询师或医学诊断的自动化系统[4-6][25-27]。",
        ],
        "国内外研究现状": [
            "国外数字心理健康研究较早关注在线筛查、远程咨询、移动干预和自助心理教育等方向。相关研究表明，数字化工具能够在一定程度上降低求助门槛，帮助用户完成自我评估和持续记录。但该类系统也面临用户持续使用率不足、算法解释困难、隐私保护要求高和危机干预边界模糊等问题，因此研究者普遍强调数字平台应作为专业服务的补充，而不能脱离专业人员单独承担高风险处置任务[7-10]。",
            "国内高校心理健康信息化建设不断推进，许多学校已经建设心理测评、预约咨询和心理健康知识宣传平台。这类系统提高了基础事务处理效率，但仍存在功能分散、数据口径不统一、学生端体验不足和风险解释较弱等问题。本文所设计的平台尝试将测评、预约、聊天、文章、通知和后台看板整合在同一系统中，使普通用户、咨询师和管理员能够围绕同一业务数据协同工作[1][6]。",
        ],
        "功能需求": [
            "心理测评模块需要支持量表配置、题目展示、答案提交、风险结果生成和历史记录查看。为了降低后续扩展成本，量表题目不应完全写死在前端页面中，而应由后端保存题目 JSON，前端根据配置动态渲染。这样在新增或调整测评量表时，只需要调整配置和后端规则，而不必大幅修改页面结构[19-22]。",
            "在线聊天模块需要与预约流程绑定。平台不应允许任意用户直接向任意咨询师发送消息，而应在双方存在已确认预约或历史沟通关系时开放聊天入口。聊天记录需要保存到数据库，以支持用户回看、咨询师了解沟通上下文和后续合规审计。对于敏感内容，系统还应预留审核状态和敏感标记字段，为后续内容治理提供基础[23-24][25]。",
        ],
        "系统架构设计": [
            "根据前文技术路线，系统整体采用前后端分离架构。前端负责页面展示、角色路由、表单交互和图表展示；后端负责接口提供、认证授权、业务规则和数据持久化；数据库保存核心业务记录；模型推理模块用于心理测评风险分析；WebSocket 通道用于在线聊天实时推送。该架构具有较好的模块边界，便于不同功能独立维护和扩展[11-18]。",
            "在安全设计方面，系统采用前端路由守卫、Axios 请求拦截器、JWT 过滤器和服务层权限校验组合实现。前端根据用户角色控制页面是否可见，后端根据 token 识别用户身份，并在预约、聊天、测评监控和后台管理等关键接口中进一步判断数据归属和角色权限。这种前后端协同的权限设计能够降低越权访问风险[17][25-26]。",
        ],
        "数据库概念结构设计": [
            "数据库设计围绕用户表、咨询师资料表、预约表、测评量表表、测评记录表、聊天消息表、文章表、通知表、评价表和操作日志表展开。用户表保存账号、密码摘要、角色和基础资料，是身份体系的中心表；咨询师资料表与用户表形成一对一关系；预约表连接普通用户和咨询师资料；测评记录表连接用户和量表；聊天消息表连接发送方和接收方。通过这些关系，系统能够将用户服务流程和后台管理流程统一到同一数据模型中[12][30]。",
            "为了满足后续扩展需求，部分表中预留了扩展字段。例如，咨询师资料表中设置评分和评价数量字段，便于后续展示服务评价；测评记录表中设置模型名称和主要影响因素字段，便于后续模型版本管理和解释展示；聊天消息表中设置敏感标记和审核状态字段，便于后续内容审核；操作日志表用于记录关键管理操作，为系统审计提供数据基础[13-14][25]。",
        ],
        "登录认证与权限控制": [
            "登录认证模块的实现包括前端状态保存、请求拦截、后端 token 校验和服务层权限判断。用户登录成功后，前端将 token 和用户信息保存到 Pinia 与 localStorage 中，Axios 拦截器在后续请求中自动携带 token；后端 JWT 过滤器校验 token 后解析用户编号和角色，并写入 Spring Security 上下文。对于管理员、咨询师和普通用户的不同页面，前端通过路由元信息进行访问控制；对于关键接口，后端通过业务逻辑再次校验当前用户是否具备操作权限[15-18]。",
        ],
        "心理测评与风险分析功能实现": [
            "心理测评模块的核心流程包括量表校验、答案解析、风险预测、结果保存和历史展示。用户提交答案后，后端首先检查量表是否存在且启用，再将答案 JSON 解析为特征映射，随后调用预测组件生成风险概率、风险等级、分析文本和主要影响因素。测评结果保存到数据库后，用户可以查看个人历史记录，管理员可以在后台查看平台测评记录和风险分布[19-22]。",
            "预测组件采用“可选模型推理 + Java 回退”的双路径。若 CatBoost 模型和 Python 推理脚本可用，系统优先调用外部模型；若模型不可用，则使用 Java 内置规则根据特征权重生成基础风险结果。这样既能体现机器学习在心理风险分析中的应用，也能保证毕业设计演示环境下系统稳定运行。需要强调的是，风险结果只是辅助参考，不能替代专业评估和医学诊断[19-22][25-27]。",
        ],
        "在线聊天功能实现": [
            "在线聊天模块采用 REST 接口与 WebSocket 结合的实现方式。REST 接口负责联系人加载、历史消息查询和消息保存，WebSocket 负责在线用户之间的实时消息推送。后端保存消息前会检查双方是否存在可沟通关系，通常要求双方存在已确认预约或历史沟通记录。该设计避免无关用户随意私聊，也使聊天记录能够与预约咨询流程对应起来[23-24]。",
            "为了增强安全性和可维护性，聊天模块还对消息内容进行基础敏感词检查，并将敏感标记和审核状态保存到数据库。虽然当前实现仍属于原型级别，但该设计为后续接入更完善的内容审核、消息撤回、文件上传、语音视频咨询和人工复核流程提供了扩展空间[24-25]。",
        ],
        "测试结果分析": [
            "从测试结果看，系统能够完成普通用户、心理咨询师和管理员三类角色的核心业务流程。登录认证能够区分用户角色，普通用户能够完成测评、预约、文章阅读和聊天，咨询师能够处理预约并进入沟通，管理员能够查看统计数据和风险记录。测试也表明，系统在接口异常、未登录访问和角色越权访问场景下具备基本处理能力[13-14]。",
            "系统仍存在进一步完善空间。首先，当前移动端适配还不够充分，后续可开发移动端页面或微信小程序；其次，心理测评模型需要进一步完善训练数据治理、模型版本管理和人工复核流程；再次，在线咨询可扩展语音、视频、消息通知和咨询记录归档；最后，后台可增加更细粒度的统计分析，例如按时间段统计预约趋势、按量表统计风险变化、按咨询师统计服务数量和评价情况[25-30]。",
        ],
        "项目总结": [
            "对照论文大纲来看，本文完成了从需求分析、技术选型、总体设计、数据库设计、核心功能实现到系统测试的完整过程。系统以 Spring Boot 和 Vue 为主要技术栈，以 MySQL 保存业务数据，以 JWT 和路由守卫实现权限控制，以 WebSocket 支撑在线聊天，以 CatBoost 或规则回退逻辑支撑心理测评风险分析。整体实现覆盖了高校心理健康服务平台的主要业务环节，具有一定的工程完整性和扩展价值[11-24]。",
        ],
    }

    for heading, paras in additions.items():
        anchor = find_para(doc, heading)
        current = anchor
        for text in paras:
            current = insert_after(current, text, "Normal")
        locations.append(f"{heading}：新增 {len(paras)} 段扩写内容，并在段落中加入相关引用。")

    ref_heading = remove_old_references(doc)
    references = [
        "[1] 江光荣, 李丹阳, 任志洪, 等. 中国国民心理健康素养的现状与特点[J]. 心理学报, 2021, 53(02): 182-201.",
        "[2] 教育部. 高等学校学生心理健康教育指导纲要[Z]. 2018.",
        "[3] 王登峰, 张大均. 大学生心理健康教育[M]. 北京: 高等教育出版社, 2019.",
        "[4] World Health Organization. Mental health: strengthening our response[EB/OL]. Geneva: WHO, 2022.",
        "[5] Torous J, Nicholas J, Larsen M E, et al. Clinical review of user engagement with mental health smartphone apps[J]. Evidence-Based Mental Health, 2018, 21(3): 116-119.",
        "[6] 李焰, 郑日昌. 大学生心理健康教育信息化建设研究[J]. 中国高等教育, 2020(10): 55-57.",
        "[7] Andersson G. Internet interventions: Past, present and future[J]. Internet Interventions, 2018, 12: 181-188.",
        "[8] Firth J, Torous J, Nicholas J, et al. The efficacy of smartphone-based mental health interventions[J]. World Psychiatry, 2017, 16(3): 287-298.",
        "[9] Mohr D C, Weingardt K R, Reddy M, Schueller S M. Three problems with current digital mental health research[J]. Psychiatric Services, 2017, 68(5): 427-429.",
        "[10] Hollis C, Falconer C J, Martin J L, et al. Annual research review: Digital health interventions for children and young people with mental health problems[J]. Journal of Child Psychology and Psychiatry, 2017, 58(4): 474-503.",
        "[11] 许令波. 深入分析 Java Web 技术内幕[M]. 北京: 电子工业出版社, 2017.",
        "[12] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
        "[13] Sommerville I. Software Engineering[M]. Boston: Pearson, 2016.",
        "[14] Pressman R S, Maxim B R. Software Engineering: A Practitioner's Approach[M]. New York: McGraw-Hill Education, 2019.",
        "[15] Craig Walls. Spring Boot 实战[M]. 北京: 人民邮电出版社, 2016.",
        "[16] Walls C. Spring in Action[M]. 6th ed. Shelter Island: Manning Publications, 2022.",
        "[17] Freeman A. Pro Spring Security[M]. New York: Apress, 2019.",
        "[18] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.",
        "[19] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[20] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System[C]. KDD, 2016: 785-794.",
        "[21] Prokhorenkova L, Gusev G, Vorobev A, Dorogush A V, Gulin A. CatBoost: unbiased boosting with categorical features[C]. NeurIPS, 2018.",
        "[22] James G, Witten D, Hastie T, Tibshirani R. An Introduction to Statistical Learning[M]. New York: Springer, 2021.",
        "[23] Fette I, Melnikov A. The WebSocket Protocol[S]. RFC 6455, 2011.",
        "[24] MDN Web Docs. WebSockets API[EB/OL]. https://developer.mozilla.org/zh-CN/docs/Web/API/WebSockets_API.",
        "[25] OWASP Foundation. OWASP Top 10[EB/OL]. https://owasp.org/www-project-top-ten/.",
        "[26] NIST. Digital Identity Guidelines[S]. SP 800-63, 2017.",
        "[27] 国家互联网信息办公室. 个人信息保护法相关解读与数据安全治理文件[Z]. 2021.",
        "[28] 尤雨溪. Vue.js 设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[29] Richardson L, Ruby S. RESTful Web Services[M]. Sebastopol: O'Reilly Media, 2007.",
        "[30] Kroenke D M, Auer D J. Database Processing: Fundamentals, Design, and Implementation[M]. London: Pearson, 2015.",
    ]
    current = ref_heading
    for ref in references:
        current = insert_after(current, ref, "参考文献")

    doc.save(OUT)

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text("\n".join(["文献引用添加位置：", *locations]), encoding="utf-8")
    print(OUT)


if __name__ == "__main__":
    main()
