from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("project_current_2_cjk20000_review.docx")
OUT = Path("project_current_2_中文字符约20000版.docx")
REPORT = Path("docs/project_current_2_reference_locations_20000cjk.txt")


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = "Normal"
    new_para.add_run(text)
    return new_para


def find_para(doc: Document, exact: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    raise ValueError(exact)


doc = Document(SRC)
anchor = find_para(doc, "系统架构设计")
anchor = insert_after(anchor, "系统架构设计还应考虑后续维护和部署便利性。对于毕业设计系统而言，开发阶段通常在本地环境运行，但论文设计不能只停留在本地演示层面，还应说明系统迁移到服务器后的基本形态。前端构建后可以作为静态资源部署，后端以独立服务形式提供接口，数据库和缓存服务作为基础设施运行。通过环境变量配置数据库连接、JWT 密钥、文件上传路径和模型脚本路径，可以减少不同环境之间的修改成本，也能避免将敏感配置直接写入代码仓库[15-18][25]。")
anchor = insert_after(anchor, "此外，平台后续若接入学校统一身份认证、短信通知、微信小程序或数据中台，也应尽量复用现有接口和数据模型。统一身份认证可以替代当前账号密码登录，短信或站内通知可以增强预约提醒能力，小程序可以提升移动端使用体验，数据中台则可以在脱敏后汇总平台运营指标。由于本系统已经采用前后端分离和模块化接口设计，后续扩展时可以围绕认证、通知、终端适配和统计接口逐步演进，而不需要推翻现有主体结构[13-14][28-30]。")
doc.save(OUT)
with REPORT.open("a", encoding="utf-8") as f:
    f.write("\n第三轮补写位置：系统架构设计，补充部署维护、统一认证、通知、小程序和数据中台扩展说明。")
print(OUT)
