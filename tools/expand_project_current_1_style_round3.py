from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("project_current_1_style_review.docx")
OUT = Path("project_current-1_扩充润色版.docx")
REPORT = Path("docs/project_current-1_修改说明.txt")


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = "Normal"
    new_para.add_run(text)
    return new_para


def find_exact(doc: Document, exact: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(exact)


doc = Document(SRC)
logs = []
more = {
    "登录认证与权限控制": [
        "登录认证与权限控制模块需要同时考虑前端体验和后端安全。前端在用户登录后保存 token，并根据角色把用户引导到相应首页，这样用户不会看到与自身身份无关的入口；后端则通过 JWT 过滤器解析请求身份，再由服务层判断具体业务是否允许操作。也就是说，前端负责让页面使用更顺畅，后端负责把真正的权限边界守住，二者配合后才能形成比较完整的访问控制链路[17-18]。",
        "在实际运行中，token 过期、用户主动退出和接口返回未授权状态都需要被统一处理。当前端发现接口返回 401 状态时，应清空本地 token 和用户信息，并跳转到登录页面，避免用户继续停留在已经失效的业务页面。对于管理员接口和咨询师处理预约接口，后端不能只相信前端传来的角色，而要根据 token 中解析出的用户编号重新查询或校验，这样可以减少伪造请求造成的风险[25-26]。",
    ],
    "咨询师端功能实现": [
        "咨询师端功能的核心是提高服务处理效率。咨询师登录后需要快速看到与自己相关的预约申请，包括预约用户、预约时间、咨询方式、问题描述和当前状态。对于待确认预约，咨询师可以根据自己的时间安排进行同意或拒绝；对于已经确认的预约，系统应提供进入聊天的入口，方便咨询师继续和用户沟通。这样的设计把预约处理和在线沟通连接在一起，有助于减少信息断裂。",
        "咨询师资料维护也是咨询师端的重要内容。职称、擅长方向、从业年限、个人介绍、咨询价格和在线状态都会影响用户选择咨询师，因此这些信息应允许咨询师或管理员进行维护。资料展示不宜过度复杂，但应包含用户做出选择所需要的关键内容；如果后续增加评价功能，还可以把评分和评价数量加入咨询师列表，使用户能够更直观地了解服务情况。",
    ],
    "兼容性与异常测试": [
        "兼容性测试主要关注系统在不同浏览器和不同窗口宽度下的表现。由于系统采用 Web 前端实现，用户可能使用 Chrome、Edge 或其他主流浏览器访问平台，也可能在较窄的屏幕中打开页面。测试时需要观察登录页、首页、表单页、表格页和聊天页是否有明显遮挡、错位或按钮不可点击的问题，尤其是聊天窗口和后台表格这类内容较多的页面，更需要注意布局稳定性。",
        "异常测试则更关注系统在不理想条件下的表现。例如，用户未登录时直接访问业务页面，系统应跳转到登录页；普通用户访问管理员页面时，应被路由守卫或后端接口拦截；提交空测评答案、无效预约时间或不存在的咨询师编号时，系统应给出明确提示；WebSocket 连接断开时，历史消息仍然应能通过接口加载。这些异常场景能反映系统是否只在理想情况下可用，还是具备一定的健壮性[13-14][25]。",
    ],
    "不足与展望": [
        "从长期建设角度看，系统还可以继续补充更细致的心理服务流程，例如知情同意、咨询记录模板、转介建议、危机工单和随访记录。当前系统已经实现了测评、预约和聊天的基本闭环，但真实高校心理服务往往需要线下制度配合，特别是面对高风险个体时，平台只能提供线索和记录，不能替代专业人员判断。因此，后续扩展应把技术功能和学校心理中心工作流程结合起来[1-6][25-27]。",
    ],
}
for heading, paras in more.items():
    current = find_exact(doc, heading)
    for text in paras:
        current = insert_after(current, text)
        logs.append(("（该位置已有基础说明，但实现细节还可以继续补充）", text, f"“{heading}”小节新增段落", "补充实现细节和测试说明，合并相关短句，使用较自然的连接方式。"))

doc.save(OUT)
with REPORT.open("a", encoding="utf-8") as f:
    for old, new, pos, reason in logs:
        f.write("\n原文：\n")
        f.write(old + "\n")
        f.write("修改后：\n")
        f.write(new + "\n")
        f.write(f"具体位置：{pos}\n")
        f.write(f"修改原因：{reason}\n")

text = "\n".join(p.text for p in doc.paragraphs)
print(OUT)
print("cjk", len(re.findall(r"[\u4e00-\u9fff]", text)))
print("nonspace", len("".join(text.split())))
