from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current-1_中英文文献重查版.docx"
OUT = ROOT / "project_current-1_中英文文献补充版.docx"

EXTRA_REFS = [
    "尚瑞莉. 互联网时代大学生心理健康教育教学模式发展探讨[J]. 中国学校卫生，2023，44(5)：801-802.",
    "李晨. 互联网时代大学生心理健康教育的突出问题与改进路径[J]. 科教文汇（中旬刊），2020(29)：155-156.",
    "辛俊杰. “互联网+”时代大学生心理健康教育的创新研究[J]. 才智，2024(30)：115-118.",
]

CITATION_REPLACEMENTS = {
    "限制[1-4]。": "限制[1-4][31-33]。",
    "意愿[1-10]。": "意愿[1-10][31-33]。",
    "策略[21-30]。": "策略[21-30][31-33]。",
}


def set_ref_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_ref_font(run)
    return p


def replace_citations(doc):
    for p in doc.paragraphs:
        text = p.text
        replaced = text
        for old, new in CITATION_REPLACEMENTS.items():
            replaced = replaced.replace(old, new)
        if replaced != text:
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)


def main():
    doc = Document(SRC)
    replace_citations(doc)

    last_ref = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("[30]"):
            last_ref = p
    if last_ref is None:
        raise RuntimeError("未找到第30条参考文献")
    cursor = last_ref
    for offset, ref in enumerate(EXTRA_REFS, start=31):
        cursor = insert_after(cursor, f"[{offset}] {ref}")

    doc.save(OUT)
    print(OUT)
    print("total references 33; chinese 23; english 10")


if __name__ == "__main__":
    main()
