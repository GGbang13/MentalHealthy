from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current-1_扩充润色版.docx"
OUT = ROOT / "project_current-1_参考文献核查替换版.docx"

NEW_REFS = [
    "江光荣，李丹阳，任志洪，等. 中国国民心理健康素养的现状与特点[J]. 心理学报，2021，53(2)：182-198.",
    "中共教育部党组. 高等学校学生心理健康教育指导纲要[Z]. 2018.",
    "李雄鹰，拜尔娜·阿里甫，檀晨唏. 学校一体化心理健康教育体系建设[J]. 中国学校卫生，2024，45(5)：615-619.",
    "靳宇倡，张政，郑佩璇，等. 远程心理健康服务：应用、优势及挑战[J]. 心理科学进展，2022，30(1)：141-156.",
    "杨晶，余林. 网络心理咨询的实践及其存在的问题[J]. 心理科学进展，2007，15(1)：140-145.",
    "姜妍，雷鸣，吴柔嘉. 高校网络心理健康教育平台建设：基于“双一流”高校的情况分析[J]. 中国心理学前沿，2020，2(7)：660-668.",
    "尚瑞莉. 互联网时代大学生心理健康教育教学模式发展探讨[J]. 中国学校卫生，2023，44(5)：801-802.",
    "王云. 互联网时代的大学生心理健康教育发展路径[J]. 青少年心理研究，2025，5(1)：9-12.",
    "刘东佳，段碧花. 大学生心理健康教育课程研究的回顾与展望：基于CiteSpace文献计量分析[J]. 中国心理学前沿，2024，6(9)：1773-1783.",
    "阮佳鑫. 网络心理咨询伦理与对策综述[J]. 中国心理学前沿，2022，4(12)：1456-1462.",
    "李海燕，吴元业，崔一鸣. 基于微服务架构的机构学者库平台设计与实现[J]. 现代信息科技，2021，5(7)：10-13，4.",
    "陈宇. 基于Spring Boot的电商管理系统的设计[J]. 现代信息科技，2020，4(1)：25-26.",
    "赵富强，严风硕，边岱泉，等. 基于Vue和SpringBoot的机场气象信息系统设计与实现[J]. 现代信息科技，2020，4(21)：1-6.",
    "杜瑛，刘冬杰. 基于Spring Boot+Vue的场地预约管理系统的设计[J]. 电脑知识与技术，2022，18(23).",
    "赵志威，张生月，蒋应举，等. 基于SpringBoot的高新技术企业创新能力评价平台设计与实现[J]. 现代信息科技，2021，5(15)：40-42.",
    "黄浩铭，刘成珏，郑滢，等. 基于Vue.js和Javalin框架的学生宿舍社交平台WebAPP应用程序的设计与开发[J]. 软件工程与应用，2019，8(6)：390-401.",
    "李玥，房科言，李茜. 基于Spring Boot的程序设计课程考试系统的设计与实现[J]. 电脑知识与技术，2025(36).",
    "韩中豪，王梓名. 基于Spring Boot框架的电子评标系统的设计与实现[J]. 现代信息科技，2019(23).",
    "朱玉琴，周碧林，张亚，等. 基于H5的钟山区商品房销售管理系统的设计与实现[J]. 现代信息科技，2022，6(18)：32-36.",
    "王丹，孙晓宇，杨路斌，等. 基于SpringBoot的软件统计分析系统设计与实现[J]. 软件工程，2019，22(3)：40-42.",
    "邓笑. 基于Spring Boot的校园轻博客系统的设计与实现[D]. 武汉：华中科技大学，2018.",
    "吕英华. 渐进式JavaScript框架Vue.js的全家桶应用[J]. 电子技术与软件工程，2019(22)：39-40.",
    "吴晶玲，简璐丝. 基于心理健康素养的大学生心理健康教育课程有效性评价标准初探[J]. 科教文汇，2020(8)：153-154.",
    "应丽莎. 线上线下有机结合的大学生心理健康教育课程教学改革[J]. 中国教育技术装备，2022(11)：105-107.",
    "张翠翠. 大学生心理健康教育课程混合式教学设计探讨：以“大学生压力管理与挫折应对”为例[J]. 教育教学论坛，2021(12)：141-144.",
    "辛俊杰. “互联网+”时代大学生心理健康教育的创新研究[J]. 才智，2024(30)：115-118.",
    "李晨. 互联网时代大学生心理健康教育的突出问题与改进路径[J]. 科教文汇，2020(29)：155-156.",
    "冷洁. “互联网+”背景下大学生心理健康教育建设路径研究[J]. 科学咨询，2024(10)：78-82.",
    "刘东江. 互联网时代数字化转型与大学生心理健康教育的融合发展[J]. 老字号品牌营销，2021(9)：155-156.",
    "陈天奇，盖斯特林. XGBoost：一种可扩展的树提升系统[C]//第22届ACM SIGKDD知识发现与数据挖掘国际会议论文集，2016：785-794.",
]


def font_run(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    font_run(run)
    return p


def main():
    doc = Document(SRC)
    ref_heading = None
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            ref_heading = p
            break
    if ref_heading is None:
        raise SystemExit("reference heading not found")
    node = ref_heading._p.getnext()
    while node is not None:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt
    cursor = ref_heading
    for idx, ref in enumerate(NEW_REFS, 1):
        cursor = insert_after(cursor, f"[{idx}] {ref}")
    doc.save(OUT)
    print(OUT)
    print(len(NEW_REFS))


if __name__ == "__main__":
    main()
