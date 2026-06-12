from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current-1_参考文献核查替换版.docx"
if not SRC.exists():
    SRC = ROOT / "project_current-1_扩充润色版.docx"
OUT = ROOT / "project_current-1_中英文文献重查版.docx"


REFERENCES = [
    "江光荣，李丹阳，任志洪，等. 中国国民心理健康素养的现状与特点[J]. 心理学报，2021，53(2)：182-198.",
    "明志君，陈祉妍. 心理健康素养：概念、评估、干预与作用[J]. 心理科学进展，2020，28(1)：1-12.",
    "靳宇倡，张政，郑佩璇，等. 远程心理健康服务：应用、优势及挑战[J]. 心理科学进展，2022，30(1)：141-156.",
    "杨晶，余林. 网络心理咨询的实践及其存在的问题[J]. 心理科学进展，2007，15(1)：140-145.",
    "何元庆，方存峰. 高校心理咨询中保密的伦理困境与解决出路[J]. 中国临床心理学杂志，2015，23(2)：378-380，377.",
    "李永慧. 高校心理危机干预中的伦理困境及应对策略[J]. 思想理论教育，2016(5)：85-88.",
    "周旻，石大维. 高校网络化心理咨询研究[J]. 中国电化教育，2015(7)：128-132.",
    "黄渊基，熊敏秀. 网络心理咨询：含义、类型及其发展[J]. 邵阳学院学报（社会科学版），2014(6)：115-120.",
    "石国兴. 英国心理咨询的专业化发展及其问题[J]. 心理科学进展，2004，12(2)：304-311.",
    "王国华. 高校心理咨询督导机制的建立与完善[J]. 黑龙江科学，2014(12)：134-135.",
    "陈宇. 基于Spring Boot的电商管理系统的设计[J]. 现代信息科技，2020，4(1)：25-26.",
    "王丹，孙晓宇，杨路斌，等. 基于SpringBoot的软件统计分析系统设计与实现[J]. 软件工程，2019，22(3)：40-42.",
    "赵富强，严风硕，边岱泉，等. 基于Vue和SpringBoot的机场气象信息系统设计与实现[J]. 现代信息科技，2020，4(21)：1-6.",
    "吕英华. 渐进式JavaScript框架Vue.js的全家桶应用[J]. 电子技术与软件工程，2019(22)：39-40.",
    "杜瑛，刘冬杰. 基于Spring Boot+Vue的场地预约管理系统的设计[J]. 电脑知识与技术，2022，18(23)：31-32.",
    "赵志威，张生月，蒋应举，等. 基于SpringBoot的高新技术企业创新能力评价平台设计与实现[J]. 现代信息科技，2021，5(15)：40-42.",
    "李海燕，吴元业，崔一鸣. 基于微服务架构的机构学者库平台设计与实现[J]. 现代信息科技，2021，5(7)：10-13，4.",
    "朱玉琴，周碧林，张亚，等. 基于H5的钟山区商品房销售管理系统的设计与实现[J]. 现代信息科技，2022，6(18)：32-36.",
    "吴晶玲，简璐丝. 基于心理健康素养的大学生心理健康教育课程有效性评价标准初探[J]. 科教文汇，2020(8)：153-154.",
    "应丽莎. 线上线下有机结合的大学生心理健康教育课程教学改革[J]. 中国教育技术装备，2022(11)：105-107.",
    "Jorm A F, Korten A E, Jacomb P A, et al. Mental health literacy: a survey of the public's ability to recognise mental disorders and their beliefs about the effectiveness of treatment[J]. Medical Journal of Australia, 1997, 166(4): 182-186.",
    "Wei Y, McGrath P J, Hayden J, Kutcher S. Mental health literacy measures evaluating knowledge, attitudes and help-seeking: a scoping review[J]. BMC Psychiatry, 2015, 15: 291.",
    "O'Connor M, Casey L, Clough B. Measuring mental health literacy: a review of scale-based measures[J]. Journal of Mental Health, 2014, 23(4): 197-204.",
    "Backhaus A, Agha Z, Maglione M L, et al. Videoconferencing psychotherapy: a systematic review[J]. Psychological Services, 2012, 9(2): 111-131.",
    "Andersson G. Internet interventions: past, present and future[J]. Internet Interventions, 2018, 12: 181-188.",
    "Torous J, Nicholas J, Larsen M E, et al. Clinical review of user engagement with mental health smartphone apps: evidence, theory and improvements[J]. Evidence-Based Mental Health, 2018, 21(3): 116-119.",
    "Fitzpatrick K K, Darcy A, Vierhile M. Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent: a randomized controlled trial[J]. JMIR Mental Health, 2017, 4(2): e19.",
    "Prokhorenkova L, Gusev G, Vorobev A, et al. CatBoost: unbiased boosting with categorical features[C]//Advances in Neural Information Processing Systems. 2018, 31: 6638-6648.",
    "Chen T, Guestrin C. XGBoost: a scalable tree boosting system[C]//Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016: 785-794.",
    "Fielding R T. Architectural styles and the design of network-based software architectures[D]. Irvine: University of California, 2000.",
]


CITATION_REPLACEMENTS = {
    "限制[1]。": "限制[1-4]。",
    "体验[15-18][23-24]。": "体验[11-18][28-30]。",
    "问题[12-14]。": "问题[11-16][30]。",
    "意愿[1-6]。": "意愿[1-10]。",
    "状态[6][13-14]。": "状态[3-8][15]。",
    "健壮性[13-14][25]。": "健壮性[11-18][23-25]。",
    "阅读[13-14]。": "阅读[11-16]。",
    "反馈[25-26]。": "反馈[21-26]。",
    "策略[25-30]。": "策略[21-30]。",
}


def set_ref_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_ref_font(run)
    return p


def rebuild_references(doc):
    ref_heading = None
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            ref_heading = p
            break
    if ref_heading is None:
        raise RuntimeError("未找到参考文献标题")
    node = ref_heading._p.getnext()
    while node is not None:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt
    cursor = ref_heading
    for idx, ref in enumerate(REFERENCES, 1):
        cursor = insert_after(cursor, f"[{idx}] {ref}")


def replace_citations(doc):
    for p in doc.paragraphs:
        text = p.text
        replaced = text
        for old, new in CITATION_REPLACEMENTS.items():
            replaced = replaced.replace(old, new)
        if replaced != text:
            # Keep the paragraph style and rebuild text; only citation paragraphs are touched.
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = replaced
            else:
                p.add_run(replaced)


def main():
    doc = Document(SRC)
    replace_citations(doc)
    rebuild_references(doc)
    doc.save(OUT)
    print(OUT)
    print("references", len(REFERENCES), "chinese", 20, "english", 10)


if __name__ == "__main__":
    main()
