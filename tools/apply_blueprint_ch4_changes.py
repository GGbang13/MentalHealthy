from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT = ROOT / "project_current-1_扩充润色版.docx"
OUTPUT = ROOT / "project_current-1_蓝本整合修改版.docx"
ASSET = ROOT / "docs" / "thesis-assets"

DOC_PARENT = None


def set_run_font(run, size=10.5, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_text(paragraph, text, size=10.5):
    paragraph.text = ""
    run = paragraph.add_run(text)
    set_run_font(run, size=size)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(21)


def add_after(element, text="", style=None, align=None, first_line=True, size=10.5):
    p = OxmlElement("w:p")
    element.addnext(p)
    para = Paragraph(p, DOC_PARENT)
    if style:
        para.style = style
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        set_run_font(run, size=size)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    if first_line:
        para.paragraph_format.first_line_indent = Pt(21)
    else:
        para.paragraph_format.first_line_indent = Pt(0)
    return para


def add_picture_after(element, image_path: Path, width_cm=15.6):
    para = add_after(element, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return para


def remove_between(doc: Document, start_p: Paragraph, end_p: Paragraph):
    body = doc._body._element
    children = list(body)
    start = children.index(start_p._p)
    end = children.index(end_p._p)
    for child in children[start + 1:end]:
        body.remove(child)


def find_heading(doc: Document, text: str):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def remove_section(doc: Document, heading_text: str, next_heading_text: str):
    start = find_heading(doc, heading_text)
    end = find_heading(doc, next_heading_text)
    if not start or not end:
        return
    body = doc._body._element
    children = list(body)
    start_i = children.index(start._p)
    end_i = children.index(end._p)
    for child in children[start_i:end_i]:
        body.remove(child)


def rebuild_41(doc: Document):
    h41 = find_heading(doc, "系统结构设计")
    h42 = find_heading(doc, "系统架构设计")
    if not h41 or not h42:
        raise RuntimeError("cannot locate 4.1/4.2 boundary")

    h41.text = "系统功能结构设计"
    h41.style = "Heading 2"
    remove_between(doc, h41, h42)

    p = add_picture_after(h41._p, ASSET / "project-system-function.png", width_cm=15.6)
    cap = add_after(p._p, "图4-1 系统功能结构图", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=10.5)
    t1 = (
        "如图4-1所示，系统功能结构以心理健康服务过程和角色使用场景为主线进行划分，整体包括账号管理、心理测评、咨询服务、预约咨询、在线沟通、文章通知、咨询师端和管理员端等模块。"
        "账号管理模块负责注册登录、资料维护和权限识别，是用户进入平台后使用其他功能的基础；心理测评模块负责量表作答、结果生成、记录查看和风险分析，能够把用户的测评行为转化为可保存、可追溯的结果数据。"
    )
    p = add_after(cap._p, t1)
    t2 = (
        "咨询服务模块主要包括咨询师列表、详情查看和评价管理，预约咨询模块主要包括提交预约、状态跟踪和预约处理，在线沟通模块负责联系人列表、消息收发和记录保存。"
        "文章通知模块承担文章浏览、通知查看和内容发布任务；咨询师端面向咨询师角色，提供资料维护、预约审核和来访沟通功能；管理员端面向管理员角色，提供用户管理、内容管理、风险监控和数据统计功能。"
        "这样处理后，原系统结构图与系统功能图中重复说明的角色入口、业务模块和管理边界被合并到同一张功能结构图中，既减少了图表重复，也能够更清楚地表现平台主要功能之间的层级关系。"
    )
    add_after(p._p, t2)


def rebuild_flow_section(doc: Document):
    flow_h = find_heading(doc, "系统业务流程设计")
    front_h = find_heading(doc, "前端工程设计")
    if not flow_h or not front_h:
        raise RuntimeError("cannot locate flow/front boundary")
    remove_between(doc, flow_h, front_h)

    p = add_picture_after(flow_h._p, ASSET / "project-system-flow.png", width_cm=15.6)
    cap = add_after(p._p, "图4-3 系统业务流程图", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    t1 = (
        "图4-3展示了平台的系统业务流程。用户进入系统后首先输入账号和密码，系统判断登录是否成功；若登录失败，则提示错误并通过折线返回账号密码输入环节，若登录成功，则继续判断用户角色。"
        "普通用户进入用户端首页后，可以选择是否进行心理测评，若选择测评则提交量表答案并生成风险等级与测评记录，随后再判断是否预约心理咨询；若不进行测评或不预约，则进入文章阅读、消息查看等普通服务流程。"
    )
    p = add_after(cap._p, t1)
    t2 = (
        "咨询师进入工作台后查看预约申请，并判断是否同意预约，同意后确认预约并开放在线聊天，拒绝后关闭该预约。管理员进入后台后查看测评与预约数据，并判断是否存在异常记录，若存在异常则进行审核处理，若不存在则直接进入用户与文章维护流程，最后查看数据面板并保存相关记录。"
        "该流程图包含登录判断、角色判断、测评判断、预约判断、咨询师审核判断和管理员异常判断，能够体现系统业务流转中的关键分支。"
    )
    p = add_after(p._p, t2)

    p = add_picture_after(p._p, ASSET / "project-data-flow.png", width_cm=15.6)
    cap = add_after(p._p, "图4-4 系统数据流图", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    t3 = (
        "图4-4进一步说明了数据在前端、后端服务、数据库和模型分析模块之间的流动过程。用户提交表单后，前端会把参数发送到后端接口，后端服务进行权限校验、参数校验和业务处理，再把结果写入数据库或返回给页面；对于心理测评数据，系统还会将答案特征传入模型分析逻辑，得到风险概率、结果等级和主要影响因素，然后把这些结果与原始测评记录一起保存。"
    )
    p = add_after(cap._p, t3)
    t4 = (
        "从整体流程看，登录和角色识别是业务入口，心理测评、预约咨询和在线聊天是用户端的核心业务，咨询师处理和管理员监管则为服务闭环提供支撑。通过流程图和数据流图配合说明，可以同时表现业务先后关系和数据处理路径，避免只描述页面跳转而忽略后端校验、模型分析和数据持久化过程。"
    )
    add_after(p._p, t4)


def update_numbering_text(doc: Document):
    replacements = {
        "系统总体架构图": "图4-2 系统总体架构图",
        "心理健康服务平台系统流程图": "图4-3 系统业务流程图",
        "系统数据流图": "图4-4 系统数据流图",
        "数据库 E-R 图": "图4-5 数据库 E-R 图",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in replacements:
            p.text = ""
            run = p.add_run(replacements[t])
            set_run_font(run, size=10.5)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t.startswith("图4.3说明"):
            set_text(p, t.replace("图4.3", "图4-3"))
        elif "图4-5进一步说明" in t:
            set_text(p, t.replace("图4-5", "图4-4"))


def main():
    global DOC_PARENT
    shutil.copy2(INPUT, OUTPUT)
    doc = Document(str(OUTPUT))
    DOC_PARENT = doc._body

    rebuild_41(doc)
    remove_section(doc, "系统功能设计", "系统业务流程设计")
    rebuild_flow_section(doc)
    update_numbering_text(doc)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
