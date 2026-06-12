from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current-1_引用细化版.docx"
OUT = ROOT / "project_current-1_段落级引用版.docx"


def set_cn_font(run, size=Pt(10.5), font="宋体"):
    run.font.name = font
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def insert_after(paragraph, text, style="Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    p.style = style
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_cn_font(run)
    return p


def set_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
        set_cn_font(p.runs[0])
    else:
        run = p.add_run(text)
        set_cn_font(run)


def main():
    doc = Document(SRC)

    for p in list(doc.paragraphs):
        text = p.text.strip()

        if text.startswith("心理健康是高校学生成长与发展的重要基础"):
            set_text(
                p,
                "心理健康是高校学生成长与发展的重要基础，近年来，学习压力、人际关系、就业竞争和家庭环境等因素不断叠加，使得一部分学生会面临焦虑、抑郁、情绪低落、睡眠障碍和适应困难等问题，高校心理健康教育与咨询工作也需要覆盖心理知识宣传、心理测评、预约咨询、危机预警、持续跟踪和管理统计等多个环节[1-4]。"
            )
            insert_after(
                p,
                "在互联网环境下，心理健康教育逐渐从单一的线下服务转向线上线下结合的服务模式，传统模式通常依赖线下预约、人工记录和单点沟通，服务流程容易受到时间、场地和人员规模的限制，因此有必要通过信息化平台提升心理健康教育与咨询服务的可达性和连续性[5-7]。"
            )

        elif text.startswith("系统架构设计需要兼顾开发便利性和后续维护"):
            set_text(
                p,
                "系统架构设计需要兼顾开发便利性和后续维护。前端构建完成后可以作为静态资源部署，后端以 Spring Boot 服务形式运行，数据库和缓存作为基础服务提供支撑，这种前后端分离与分层设计方式能够提升系统开发、部署和维护的便利性[8-15]。"
            )
            p2 = insert_after(
                p,
                "心理测评与风险分析模块需要对测评结果进行结构化处理，并在结果解释、风险概率和影响因素记录等方面保留可追溯数据，机器学习相关研究为风险识别和模型分析提供了技术参考[16-17]。"
            )
            insert_after(
                p2,
                "浏览器端通过统一接口访问系统，普通业务走 REST API，聊天业务在保存消息后通过 WebSocket 推送，这样既能保证数据留痕，也能提升在线沟通的实时性和接口组织的规范性[18-19]。"
            )

        elif text.startswith("后端登录流程由认证控制器接收用户名和密码"):
            set_text(
                p,
                "后端登录流程由认证控制器接收用户名和密码，业务层根据用户名查询用户记录，并使用 BCrypt 对密码摘要进行校验。校验通过后，系统根据用户编号和角色生成 JWT，前端获得 token 后写入 Pinia 状态和 localStorage。该流程将会话状态放在客户端保存，服务端接口通过过滤器解析请求头中的身份令牌，从而避免传统 Session 在前后端分离场景中的状态同步问题[20-22]。"
            )

        elif text.startswith("用户端功能的实现重点是把常用入口集中起来"):
            set_text(
                p,
                "用户端功能的实现重点是把常用入口集中起来。用户登录后可以从首页进入心理测评、预约咨询、心理文章和在线沟通，页面不需要让用户在复杂菜单中反复查找。对于心理健康服务平台而言，清晰入口本身就是降低求助成本的一部分，尤其是当用户处在压力较大或情绪低落状态时，过于复杂的操作流程会降低使用意愿[24-25]。"
            )

        elif text.startswith("在预约咨询功能中，用户需要根据咨询师资料做出选择"):
            set_text(
                p,
                "在预约咨询功能中，用户需要根据咨询师资料做出选择，因此页面会展示咨询师姓名、职称、擅长方向、从业年限、咨询价格、在线状态和评分等信息。用户提交预约时填写预约时间、咨询形式、咨询时长和问题描述，系统保存后进入待确认状态，等咨询师处理后再显示后续状态[26-31]。"
            )

        elif text.startswith("异常测试则更关注系统在不理想条件下的表现"):
            set_text(
                p,
                "异常测试则更关注系统在不理想条件下的表现。例如，用户未登录时直接访问业务页面，系统应跳转到登录页；普通用户访问管理员页面时，应被路由守卫或后端接口拦截；提交空测评答案、无效预约时间或不存在的咨询师编号时，系统应给出明确提示；WebSocket 连接断开时，历史消息仍然应能通过接口加载。这些异常场景能反映系统是否只在理想情况下可用，还是具备一定的健壮性[32]。"
            )

        elif text.startswith("从功能覆盖角度看"):
            set_text(
                p,
                "从功能覆盖角度看，系统测试围绕三类角色和多条主流程展开，包括登录认证、心理测评、预约咨询、在线聊天、内容管理和数据看板。测试时不仅要验证按钮是否能够点击，还要验证前后端数据是否一致，例如预约确认后用户端是否出现聊天入口，测评提交后管理员端是否能够查看记录，文章发布后用户端是否能够阅读[33]。"
            )

        elif text.startswith("从异常处理角度看"):
            set_text(
                p,
                "从异常处理角度看，测试需要关注未登录访问、角色越权、参数缺失、数据不存在和 WebSocket 断开等情况。普通用户访问管理员页面应被拦截，咨询师处理不属于自己的预约应被拒绝，聊天双方没有确认预约关系时不能发送消息。通过这些测试，可以验证系统在边界场景下是否仍能给出明确反馈[34]。"
            )

        elif text.startswith("后续工作可以从移动端体验"):
            set_text(
                p,
                "后续工作可以从移动端体验、模型治理和平台安全三个方向继续展开。一方面，学生使用心理服务时经常依赖手机，因此可以开发响应式移动页面或微信小程序，进一步降低学生在求助过程中的使用门槛[35]。"
            )
            p2 = insert_after(
                p,
                "另一方面，心理风险分析需要更完整的数据来源说明、模型版本记录和人工复核流程，平台后续可以围绕心理健康素养、风险识别和数字化干预形成更加清晰的服务闭环[36-37]。"
            )
            insert_after(
                p2,
                "同时，平台还应补充操作日志审计、敏感数据访问审批、数据留存策略以及智能化咨询辅助能力，使心理服务平台在扩展功能时仍能兼顾伦理边界、数据安全和服务质量[38]。"
            )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
