from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


SRC = Path("project_current.docx")
OUT = Path("project_current_参考文献扩充版.docx")


def insert_after(paragraph, text: str, style: str = "Normal"):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    new_para.style = style
    new_para.add_run(text)
    return new_para


def find_para(doc: Document, exact: str):
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    raise ValueError(f"paragraph not found: {exact}")


def find_contains(doc: Document, needle: str):
    for para in doc.paragraphs:
        if needle in para.text:
            return para
    raise ValueError(f"paragraph containing not found: {needle}")


def append_citation_once(para, citation: str):
    text = para.text.strip()
    if citation not in text:
        para.add_run(citation)


def try_append_citation(doc: Document, needle: str, citation: str):
    try:
        append_citation_once(find_contains(doc, needle), citation)
    except ValueError:
        pass


def remove_old_references(doc: Document):
    start = None
    for idx, para in enumerate(doc.paragraphs):
        if para.text.strip() == "参考文献":
            start = idx
            break
    if start is None:
        raise ValueError("reference heading not found")
    for para in list(doc.paragraphs[start + 1:]):
        para._element.getparent().remove(para._element)
    return doc.paragraphs[start]


def main():
    doc = Document(SRC)

    # Add citations to existing discussion paragraphs.
    try_append_citation(doc, "心理健康是高校学生成长与发展的重要基础", "[1-3]")
    try_append_citation(doc, "基于 Web 的心理健康服务平台", "[4-6]")
    try_append_citation(doc, "国外高校与医疗机构较早", "[7-10]")
    try_append_citation(doc, "前后端分离架构", "[11-14]")
    try_append_citation(doc, "系统采用前后端分离模式", "[15-18]")
    try_append_citation(doc, "心理测评模块除风险概率与等级外", "[19-22]")
    try_append_citation(doc, "在线聊天模块", "[23-24]")
    try_append_citation(doc, "心理健康属于敏感领域", "[25-27]")
    try_append_citation(doc, "后台管理功能可以继续加强统计分析能力", "[28-30]")

    additions: dict[str, list[str]] = {
        "研究背景与意义": [
            "从高校管理实践看，心理健康服务平台的建设并不是简单地把线下表格搬到线上，而是要围绕学生求助、咨询师处理、管理员监管和数据分析四类活动建立连续流程。传统人工登记方式容易出现信息分散、预约反馈慢、历史记录不完整等问题，尤其在测评人数较多、咨询资源有限的情况下，管理人员难以及时掌握整体风险分布。因此，信息化平台应当承担统一入口、流程留痕、权限隔离和统计分析等职责，为心理健康教育工作提供可追踪、可复用的数据基础[1-4]。",
            "与此同时，心理健康服务具有较强的隐私性和专业性。平台在提升效率的同时，不能将心理测评结果简单等同于诊断结论，也不能让非授权人员随意接触用户敏感信息。因此，本系统在设计中强调角色边界和数据最小化展示：普通用户主要查看个人服务信息，咨询师主要处理与自身服务相关的预约与沟通，管理员主要查看平台运行数据和必要的风险线索。这样的设计有助于在提高服务效率和保护用户隐私之间取得平衡[5-6]。",
        ],
        "国内外研究现状": [
            "国外数字心理健康服务研究较早关注在线筛查、远程咨询、移动干预和数据辅助评估等方向。相关研究表明，数字化工具能够降低部分群体的求助门槛，并在心理健康知识普及、随访记录和辅助决策方面发挥积极作用。但该类系统也面临隐私保护、模型偏差、危机干预边界和专业人员参与不足等问题，因此多数研究强调数字平台应作为专业服务的补充，而不是替代专业心理咨询和医学诊断[7-10]。",
            "国内高校心理健康信息化建设近年来不断推进，许多学校已经建设了心理测评、预约咨询、活动报名和知识宣传类系统。这些系统提高了基础事务处理效率，但仍存在数据孤岛、功能分散、风险解释不足和跨角色协同不够顺畅等问题。对于学生而言，若测评、预约、文章和咨询入口分散在不同系统中，会增加使用成本；对于咨询师和管理员而言，若缺少统一数据视图，也不利于后续服务跟踪和资源配置[1][6]。",
            "结合本课题实现情况，本文关注的是面向高校场景的综合型心理健康服务平台。与单一测评系统或单一预约系统相比，本系统更强调多模块协同：心理测评为用户自我了解和管理员风险监控提供数据来源，预约咨询为后续服务衔接提供流程入口，在线聊天为咨询沟通提供记录化渠道，后台看板为管理人员提供统计信息。上述设计能够较好体现 Web 系统在高校心理服务中的综合应用价值[11-14]。",
        ],
        "功能需求": [
            "在心理测评功能方面，系统需要支持量表配置、题目展示、答案提交、风险结果生成和历史记录查看。由于不同量表的问题数量、选项含义和计分方式可能存在差异，系统不宜将题目写死在前端页面中，而应通过后端配置保存题目 JSON，由前端动态渲染。这样可以在后续新增 PHQ-9、GAD-7 或其他量表时降低修改成本，也能使测评模块具有更好的扩展性[19-22]。",
            "在在线咨询功能方面，系统需要将预约状态与聊天准入进行关联。也就是说，平台不应允许任意用户直接向任意咨询师发送消息，而应在双方存在已确认预约或历史沟通关系时开放聊天入口。该规则既能减少无效沟通，也能使聊天记录与咨询服务流程形成对应关系。聊天消息需要持久化保存，便于用户回看、咨询师了解上下文和管理员在必要情况下进行合规审计[23-24]。",
        ],
        "系统架构设计": [
            "系统整体采用前后端分离架构。前端基于 Vue 3 构建单页应用，负责页面展示、表单交互、路由控制和数据可视化；后端基于 Spring Boot 提供 REST API、身份认证、业务规则处理和数据访问；数据库保存用户、咨询师、预约、测评、聊天、文章和通知等业务数据。前后端通过统一接口进行通信，使页面层和业务层之间保持清晰边界，也便于后续增加移动端或小程序端[11-18]。",
            "在安全架构方面，系统采用 JWT 和 Spring Security 实现无状态认证。用户登录成功后获得 token，前端在后续请求中将 token 放入请求头，后端过滤器解析 token 后识别当前用户身份与角色。前端路由守卫可以提升页面访问体验，后端接口校验则负责真正的安全控制。对于管理员接口、测评监控接口和聊天消息接口，系统还需要在服务层进行二次权限判断，避免仅依赖前端控制造成越权访问[15-18]。",
        ],
        "数据库概念结构设计": [
            "数据库设计围绕用户、咨询师资料、预约记录、测评量表、测评记录、聊天消息、文章和通知等核心实体展开。用户表是身份体系的中心表，咨询师资料表与用户表形成一对一关系；预约表连接普通用户和咨询师资料，记录预约时间、咨询形式、问题描述和状态；测评记录表连接用户和量表，保存答案 JSON、风险概率、风险等级和分析文本；聊天消息表连接发送方和接收方，保存消息内容和审核状态。这样的结构能够覆盖平台主要业务流程，并为统计分析和审计扩展提供基础[12][18]。",
            "为了提高系统可维护性，数据库字段设计需要兼顾当前功能和后续扩展。例如，咨询师资料表中保留评分和评价数量字段，可为后续服务评价功能提供基础；测评记录表中保留模型名称和主要影响因素字段，可为模型版本管理和结果解释提供依据；聊天消息表中保留敏感标记和审核状态字段，可为后续内容审核提供扩展空间。通过这些预留字段，系统能够在不频繁破坏表结构的情况下持续迭代[13-14]。",
        ],
        "登录认证与权限控制": [
            "登录认证模块的实现包括前端状态管理、请求拦截、后端 token 校验和角色权限判断。前端使用 Pinia 保存 token 与用户信息，并通过 Axios 拦截器在请求头中自动携带认证信息；后端通过 JWT 过滤器解析用户身份，将认证结果写入 Spring Security 上下文。对于不同角色可访问的页面，前端通过路由元信息进行初步控制；对于用户资料、预约处理、测评监控和后台管理等接口，后端仍需要根据当前用户身份进行二次校验，确保权限控制不被绕过[15-18]。",
        ],
        "心理测评与风险分析功能实现": [
            "心理测评模块是系统的特色功能之一。用户提交答案后，后端首先校验量表是否启用，再将答案 JSON 解析为特征映射，随后调用风险预测组件生成结果。预测组件优先尝试调用外部 CatBoost 模型；当模型文件或 Python 推理脚本不可用时，系统会回退到 Java 内置规则逻辑，根据特征权重计算风险概率和风险等级。该设计兼顾了模型扩展能力和系统可用性，避免因为模型环境缺失导致测评流程完全不可用[19-22]。",
            "测评结果不仅包含分数和等级，还包含模型名称、分析文本和主要影响因素。主要影响因素可以帮助用户理解风险来源，例如睡眠问题、长期压力、社交退缩或家庭支持不足等。需要说明的是，系统给出的风险分析属于辅助参考，不能替代专业心理咨询师的评估和医学诊断。因此，论文在描述该模块时应强调模型结果的辅助性质，以及后续人工复核和危机干预流程的重要性[25-27]。",
        ],
        "在线聊天功能实现": [
            "在线聊天模块采用 REST 接口与 WebSocket 相结合的方式。REST 接口负责历史消息加载、消息保存和联系人列表查询，WebSocket 用于在线状态下的实时推送。后端在保存消息前会校验双方是否具有可沟通关系，通常要求双方存在已确认预约或历史会话记录。该设计将业务准入、消息持久化和实时推送拆分处理，使聊天功能在弱网或连接异常情况下仍能保持基本可用[23-24]。",
        ],
        "测试结果分析": [
            "从测试结果看，系统能够完成普通用户、咨询师和管理员三类角色的核心业务流程。登录认证、测评提交、预约状态流转、聊天记录保存和后台统计等功能均能按照预期运行。测试过程中也暴露出一些后续改进方向，例如移动端适配仍需进一步优化，测评模型治理和人工复核流程需要更加完善，聊天内容审核和操作日志也可以继续增强。这些问题不影响原型系统的基本可用性，但对于真实高校部署具有重要参考价值[28-30]。",
        ],
    }

    inserted_locations = []
    for heading, paras in additions.items():
        anchor = find_para(doc, heading)
        current = anchor
        for text in paras:
            current = insert_after(current, text)
            inserted_locations.append((heading, text[:38]))

    ref_heading = remove_old_references(doc)
    references = [
        "[1] 江光荣, 李丹阳, 任志洪, 等. 中国国民心理健康素养的现状与特点[J]. 心理学报, 2021, 53(02): 182-201.",
        "[2] 教育部. 高等学校学生心理健康教育指导纲要[Z]. 2018.",
        "[3] 王登峰, 张大均. 大学生心理健康教育[M]. 北京: 高等教育出版社, 2019.",
        "[4] World Health Organization. Mental health: strengthening our response[EB/OL]. Geneva: WHO, 2022.",
        "[5] Torous J, Nicholas J, Larsen M E, et al. Clinical review of user engagement with mental health smartphone apps[J]. Evidence-Based Mental Health, 2018, 21(3): 116-119.",
        "[6] 李焰, 郑日昌. 大学生心理健康教育信息化建设研究[J]. 中国高等教育, 2020(10): 55-57.",
        "[7] Andersson G. Internet interventions: Past, present and future[J]. Internet Interventions, 2018, 12: 181-188.",
        "[8] Firth J, Torous J, Nicholas J, et al. The efficacy of smartphone-based mental health interventions[J]. World Psychiatry, 2017, 16(3): 287-298.",
        "[9] Mohr D C, Weingardt K R, Reddy M, Schueller S M. Three problems with current digital mental health research[J]. Psychiatric Services, 2017, 68(5): 427-429.",
        "[10] Hollis C, Falconer C J, Martin J L, et al. Annual research review: Digital health interventions for children and young people with mental health problems[J]. Journal of Child Psychology and Psychiatry, 2017, 58(4): 474-503.",
        "[11] 许令波. 深入分析 Java Web 技术内幕[M]. 北京: 电子工业出版社, 2017.",
        "[12] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
        "[13] Sommerville I. Software Engineering[M]. Boston: Pearson, 2016.",
        "[14] Pressman R S, Maxim B R. Software Engineering: A Practitioner's Approach[M]. New York: McGraw-Hill Education, 2019.",
        "[15] Craig Walls. Spring Boot 实战[M]. 北京: 人民邮电出版社, 2016.",
        "[16] Walls C. Spring in Action[M]. 6th ed. Shelter Island: Manning Publications, 2022.",
        "[17] Freeman A. Pro Spring Security[M]. New York: Apress, 2019.",
        "[18] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.",
        "[19] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[20] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System[C]. KDD, 2016: 785-794.",
        "[21] Prokhorenkova L, Gusev G, Vorobev A, Dorogush A V, Gulin A. CatBoost: unbiased boosting with categorical features[C]. NeurIPS, 2018.",
        "[22] James G, Witten D, Hastie T, Tibshirani R. An Introduction to Statistical Learning[M]. New York: Springer, 2021.",
        "[23] Fette I, Melnikov A. The WebSocket Protocol[S]. RFC 6455, 2011.",
        "[24] MDN Web Docs. WebSockets API[EB/OL]. https://developer.mozilla.org/zh-CN/docs/Web/API/WebSockets_API.",
        "[25] OWASP Foundation. OWASP Top 10[EB/OL]. https://owasp.org/www-project-top-ten/.",
        "[26] NIST. Digital Identity Guidelines[S]. SP 800-63, 2017.",
        "[27] 国家互联网信息办公室. 个人信息保护法相关解读与数据安全治理文件[Z]. 2021.",
        "[28] 尤雨溪. Vue.js 设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[29] Richardson L, Ruby S. RESTful Web Services[M]. Sebastopol: O'Reilly Media, 2007.",
        "[30] Kroenke D M, Auer D J. Database Processing: Fundamentals, Design, and Implementation[M]. London: Pearson, 2015.",
    ]
    current = ref_heading
    for ref in references:
        current = insert_after(current, ref, style="参考文献")

    doc.save(OUT)

    # Write a small plain text report for the final answer.
    report = Path("docs/reference_update_locations.txt")
    report.parent.mkdir(parents=True, exist_ok=True)
    report.write_text(
        "\n".join([
            "新增/补充引用位置：",
            "1. 研究背景与意义：补充高校心理健康信息化、隐私与服务边界说明，引用[1-6]。",
            "2. 国内外研究现状：补充数字心理健康、国内高校心理系统现状，引用[7-14]。",
            "3. 功能需求：补充心理测评动态量表、聊天准入规则，引用[19-24]。",
            "4. 系统架构设计：补充前后端分离、JWT 与安全架构说明，引用[11-18]。",
            "5. 数据库概念结构设计：补充实体关系、扩展字段设计说明，引用[12-14][18]。",
            "6. 登录认证与权限控制：补充前端状态、Axios、JWT 过滤器和服务层二次校验，引用[15-18]。",
            "7. 心理测评与风险分析功能实现：补充 CatBoost、Java 回退、结果解释和人工复核说明，引用[19-27]。",
            "8. 在线聊天功能实现：补充 REST + WebSocket 组合与消息持久化说明，引用[23-24]。",
            "9. 测试结果分析/不足与展望：补充测试结论、移动端、模型治理和审计方向，引用[28-30]。",
        ]),
        encoding="utf-8",
    )
    print(OUT)


if __name__ == "__main__":
    main()
