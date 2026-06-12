from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current_1_final_source.docx"
OUT_TMP = ROOT / "project_current_1_final_review.docx"
OUT_CN = ROOT / "project_current-1_最终完善版.docx"
REPORT = ROOT / "docs" / "project_current-1_最终修改说明.txt"
ASSETS = ROOT / "docs" / "thesis-assets"


REFERENCES = [
    "江光荣，李丹阳，任志洪，等. 中国国民心理健康素养的现状与特点[J]. 心理学报，2021，53(02)：182-201.",
    "教育部. 高等学校学生心理健康教育指导纲要[Z]. 2018.",
    "王登峰，张大均. 大学生心理健康教育[M]. 北京：高等教育出版社，2020.",
    "World Health Organization. Mental health: strengthening our response[R]. Geneva: WHO, 2022.",
    "Torous J, Nicholas J, Larsen M E, et al. Clinical review of user engagement with mental health smartphone apps[J]. Evidence-Based Mental Health, 2018, 21(3): 116-119.",
    "李焰，郑日昌. 大学生心理健康教育信息化建设研究[J]. 中国学校卫生，2020，41(10)：55-57.",
    "Andersson G. Internet interventions: Past, present and future[J]. Internet Interventions, 2018, 12: 181-188.",
    "Fitzpatrick K K, Darcy A, Vierhile M. Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent[J]. JMIR Mental Health, 2017, 4(2): e19.",
    "姚本先. 学校心理健康教育新论[M]. 北京：高等教育出版社，2019.",
    "周宗奎，孙晓军. 网络心理健康服务的发展与反思[J]. 心理科学进展，2019，27(08)：1362-1372.",
    "Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.",
    "Walls C. Spring Boot in Action[M]. New York: Manning Publications, 2016.",
    "Johnson R, Hoeller J, Arendsen A, et al. Professional Java Development with the Spring Framework[M]. Indianapolis: Wiley Publishing, 2005.",
    "Spilca L. Spring Security in Action[M]. New York: Manning Publications, 2020.",
    "Chodorow K. MongoDB: The Definitive Guide[M]. Sebastopol: O'Reilly Media, 2013.",
    "Beighley L, Morrison M. Head First SQL[M]. Sebastopol: O'Reilly Media, 2007.",
    "Ben-Gan I. T-SQL Fundamentals[M]. Redmond: Microsoft Press, 2016.",
    "Fowler M. Patterns of Enterprise Application Architecture[M]. Boston: Addison-Wesley, 2002.",
    "Gamma E, Helm R, Johnson R, Vlissides J. Design Patterns: Elements of Reusable Object-Oriented Software[M]. Boston: Addison-Wesley, 1995.",
    "Pressman R S, Maxim B R. Software Engineering: A Practitioner's Approach[M]. New York: McGraw-Hill Education, 2020.",
    "Sommerville I. Software Engineering[M]. Boston: Pearson, 2016.",
    "Beck K. Test Driven Development: By Example[M]. Boston: Addison-Wesley, 2003.",
    "Myers G J, Sandler C, Badgett T. The Art of Software Testing[M]. Hoboken: Wiley, 2011.",
    "Nielsen J. Usability Engineering[M]. San Francisco: Morgan Kaufmann, 1994.",
    "Wieruch R. The Road to React[M]. Leanpub, 2022.",
    "Vue.js Team. Vue.js Documentation[EB/OL]. https://vuejs.org/.",
    "MyBatis-Plus Team. MyBatis-Plus Documentation[EB/OL]. https://baomidou.com/.",
    "Auth0. Introduction to JSON Web Tokens[EB/OL]. https://jwt.io/introduction.",
    "Prokhorenkova L, Gusev G, Vorobev A, et al. CatBoost: unbiased boosting with categorical features[J]. Advances in Neural Information Processing Systems, 2018, 31: 6638-6648.",
    "Chen T, Guestrin C. XGBoost: A scalable tree boosting system[C]//Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. 2016: 785-794.",
]


SECTION_INTROS = {
    "绪论": "本章围绕心理健康服务平台的研究起点展开说明，主要从研究背景、研究意义、国内外发展情况、论文内容安排以及技术路线等方面进行梳理。通过这一部分的分析，可以把系统建设与高校心理健康服务的现实需求联系起来，也能说明这项研究为什么会选择 Spring Boot、Vue 3、MySQL、JWT、WebSocket 和 CatBoost 等技术作为主要支撑。",
    "相关技术介绍": "本章对系统开发中会用到的关键技术进行说明，内容包括后端服务框架、前端交互框架、数据持久化、安全认证、实时通信、数据可视化以及心理风险分析模型等。各项技术并不是孤立使用，而是共同服务于用户登录、心理测评、预约咨询、在线沟通和后台管理等业务流程，因此在介绍时会重点说明它们在本项目当中的具体作用。",
    "系统需求分析": "本章根据平台的实际使用场景，对系统的可行性、角色需求、功能需求和非功能需求进行分析。需求分析的重点并不是单纯罗列功能，而是把学生、咨询师和管理员三类角色在系统中的操作路径说明清楚，从而为后续数据库设计、接口设计和功能实现提供相对稳定的根据。",
    "系统总体设计": "本章在需求分析的基础上，对系统的整体结构、技术架构、功能划分、业务流程、前端工程组织、数据库模型以及接口交互方式进行设计。总体设计部分承担着承上启下的作用，一方面需要把前文提出的需求落到可实现的模块当中，另一方面也要为后文的核心功能实现说明模块之间的数据传递关系。",
    "系统核心功能实现": "本章围绕平台已经实现的核心功能展开说明，主要包括登录认证、用户端服务、咨询师端服务、管理员端服务、心理测评、在线聊天、预约咨询以及数据看板等内容。为了让实现过程更加具体，本章会结合部分关键代码片段解释业务逻辑、权限控制、状态变化和数据保存方式。",
    "系统测试": "本章主要说明平台的测试环境、测试方法、功能测试、兼容性测试和异常测试结果。测试工作的目标是验证系统能否按照需求稳定运行，也要检查权限控制、输入校验、业务状态流转和页面交互是否存在明显问题，从而提高系统交付时的可靠性。",
    "总结与展望": "本章对平台的设计与实现工作进行总结，并结合系统目前已经完成的功能说明后续还可以继续改进的方向。总结部分关注已经达成的建设成果，展望部分则从算法优化、服务闭环、数据安全和多端适配等角度提出进一步完善的可能路径。",
}

EXTRA_SECTIONS = {
    "用例模型分析": [
        "系统用例模型主要围绕学生用户、咨询师用户和管理员用户三类角色展开，学生用户能够完成注册登录、个人资料维护、心理测评、测评记录查看、预约咨询、在线聊天和查看公告等操作；咨询师用户可以维护个人简介、查看预约申请、确认或拒绝预约、与已经建立咨询关系的学生进行沟通，并在必要时关注高风险测评记录；管理员用户则负责用户管理、咨询师信息维护、测评量表维护、预约与聊天记录监管、敏感内容审核以及平台数据统计等工作。",
        "从用例之间的关系看，登录认证是大部分业务用例的前置条件，心理测评会产生风险等级和测评记录，预约咨询会进一步影响聊天权限，后台管理则会对量表、用户、预约和聊天内容进行维护与监督。这样的用例划分能把系统中的操作边界表达清楚，也能避免功能实现时出现角色权限混乱的问题。"
    ],
    "系统功能设计": [
        "系统功能设计按照“用户服务、咨询服务、测评服务、沟通服务、后台管理”几个方向进行划分。用户服务主要处理账号注册、登录认证、个人资料和角色信息；咨询服务负责咨询师信息展示、预约申请、预约审核和预约状态维护；测评服务负责量表展示、答案提交、分值计算、风险预测和历史记录查询；沟通服务负责聊天联系人生成、消息发送、消息记录查询和敏感词处理；后台管理则负责用户、咨询师、量表、预约、聊天记录和统计数据的统一维护。",
        "这种功能拆分方式一方面贴合心理健康服务平台的业务过程，另一方面也便于前后端分层开发。前端可以按照页面和用户角色组织组件，后端可以按照业务模块划分 Controller、Service、Mapper 和实体对象，数据库表也能对应到具体业务对象，从而使系统结构更加清晰。"
    ],
    "测试方法": [
        "系统测试采用功能测试、接口测试、异常测试和兼容性测试相结合的方式进行。功能测试重点检查注册登录、测评提交、预约申请、预约审核、聊天发送和后台管理等主要流程是否能够完整执行；接口测试重点验证请求参数、响应数据、身份令牌和权限拦截是否符合设计要求；异常测试主要关注空值、非法状态、无权限访问、重复提交和数据库记录缺失等情况；兼容性测试则检查不同浏览器窗口宽度下页面展示和交互是否正常。",
        "在测试过程中，笔者按照普通用户、咨询师和管理员三种身份分别登录系统，并根据业务流程逐项执行操作。对于需要状态变化的功能，例如预约申请、咨询师确认、聊天权限开启等，会先构造前置数据，再观察数据库记录和页面展示是否同步变化；对于涉及安全控制的功能，则会重点检查未登录访问、角色越权访问和无效 token 访问是否被系统拦截。"
    ],
}

FIGURES = {
    "系统结构设计": [
        ("project-system-structure.png", "图4-1 系统结构图", "如图4-1所示，系统结构按照表现层、业务层、数据访问层和数据存储层进行划分。表现层由 Vue 3 页面和组件负责呈现，业务层由 Spring Boot 中的 Controller 与 Service 承担核心处理，数据访问层通过 MyBatis-Plus Mapper 完成数据库操作，底层则由 MySQL 保存用户、咨询师、预约、测评和聊天等业务数据。这样的结构能够把页面交互、业务规则和数据持久化分开，使得后期维护时可以按照模块定位问题。"),
    ],
    "系统架构设计": [
        ("project-system-architecture.png", "图4-2 系统架构图", "图4-2展示了系统的前后端分离架构。浏览器端通过 Axios 调用后端 REST 接口，并在需要实时沟通时使用 WebSocket 完成消息交互；后端使用 Spring Security 与 JWT 进行身份校验，把通过校验的用户身份写入安全上下文，再由业务服务完成预约、测评、聊天和后台管理等处理；数据库层保存结构化业务数据，模型分析模块则为测评结果提供风险概率和等级判断。"),
    ],
    "系统功能设计": [
        ("project-system-function.png", "图4-3 系统功能图", "图4-3按照角色和业务模块对系统功能进行了拆分。学生端侧重于心理测评、咨询预约和在线沟通，咨询师端侧重于预约处理、资料维护和咨询沟通，管理员端侧重于平台维护、数据监管和统计分析。不同角色共享登录认证和基础资料管理能力，但在业务权限上有明显区分，这样可以保证平台既能服务学生，也能支撑咨询师与管理员开展工作。"),
    ],
    "系统流程设计": [
        ("project-system-flow.png", "图4-4 系统流程图", "图4-4说明了用户进入系统后的主要业务流程。用户完成登录后，系统会按照角色进入不同页面；学生可以选择测评或预约咨询，测评完成后生成风险结果，预约通过后可以与咨询师沟通；咨询师对预约申请进行处理，并在预约确认后开展交流；管理员则通过后台查看整体数据和异常内容。该流程体现了测评、预约、沟通与管理之间的衔接关系。"),
        ("project-data-flow.png", "图4-5 系统数据流图", "图4-5进一步说明了数据在前端、后端服务、数据库和模型分析模块之间的流动过程。用户提交表单后，前端会把参数发送到后端接口，后端服务进行权限校验、参数校验和业务处理，再把结果写入数据库或返回给页面；对于心理测评数据，系统还会将答案特征传入模型分析逻辑，得到风险概率、结果等级和主要影响因素，然后把这些结果与原始测评记录一起保存。"),
    ],
    "用例模型分析": [
        ("project-use-case-model.png", "图3-1 系统用例模型图", "图3-1从参与者角度展示了系统用例关系。学生、咨询师和管理员分别承担不同操作，登录认证作为多个用例的共同前置条件，预约咨询又会影响在线聊天的可用性。通过用例模型可以较直观地看到系统边界、角色权限和功能依赖关系，为后续的模块设计提供了依据。"),
    ],
    "数据库设计": [
        ("project-er-diagram-detailed.png", "图4-6 数据库ER图", "图4-6展示了系统主要实体之间的关系。用户表是系统的基础表，咨询师资料表与用户表通过 user_id 形成关联，预约表连接学生用户与咨询师资料，测评记录表关联用户和测评量表，聊天消息表则通过发送者和接收者记录沟通关系。这样的实体关系能够支撑平台完成身份区分、测评归档、预约跟踪和聊天追溯。"),
        ("project-database-table-design.png", "图4-7 数据库表结构设计图", "图4-7对主要数据表及字段进行了整理。表结构设计时既考虑了业务字段，例如预约时间、预约状态、测评分数、风险概率、聊天内容等，也保留了创建时间、更新时间、状态字段等维护信息，这样可以支持后续的数据统计、记录筛选和异常追踪。"),
    ],
    "后端接口与数据交互设计": [
        ("project-interface-design-table.png", "图4-8 后端接口设计表", "图4-8列出了系统中的主要接口。接口设计按照业务模块划分路径，例如认证接口、测评接口、预约接口、聊天接口和后台管理接口等；前端在调用接口时会携带 token，后端根据 token 中的用户编号和角色判断是否允许访问，从而保证接口数据交互过程具有明确的权限边界。"),
    ],
    "心理测评与风险分析功能实现": [
        ("project-assessment-flow-detail.png", "图5-1 心理测评与风险分析流程图", "图5-1说明了心理测评功能的实现过程。用户选择量表并提交答案后，后端会先检查量表是否有效，再把答案 JSON 解析为模型可以识别的变量特征，随后调用预测逻辑得到分数、风险概率、结果等级、分析文本和主要影响因素，最后将完整记录写入数据库，供用户端历史记录和管理员端风险监测使用。"),
    ],
    "在线聊天功能实现": [
        ("project-chat-flow.png", "图5-2 在线聊天功能流程图", "图5-2展示了在线聊天模块的数据处理过程。系统在保存消息前会先判断双方是否存在可沟通关系，只有已经确认的预约关系或已有沟通记录才允许继续发送；消息保存时会进行敏感词检测和脱敏处理，并根据检测结果设置审核状态。这样做能把咨询沟通与预约流程联系起来，也能减少不当内容直接展示造成的风险。"),
    ],
    "预约咨询功能实现": [
        ("project-appointment-flow.png", "图5-3 预约咨询流程图", "图5-3展示了预约咨询从申请到处理的状态变化。学生提交预约后，记录进入 PENDING 状态；咨询师确认后，预约状态变为 CONFIRMED，同时提醒状态变为 READY，系统会开放聊天入口；如果咨询师拒绝预约，则状态变为 REJECTED，提醒状态关闭。该设计能让预约、提醒和聊天权限保持一致。"),
    ],
    "登录认证与权限控制": [
        ("project-auth-flow.png", "图5-4 登录认证流程图", "图5-4说明了系统的认证流程。用户登录成功后，后端签发 JWT，前端在后续请求中把 token 放入 Authorization 请求头；后端过滤器解析并校验 token，校验通过后把用户编号和角色写入 Spring Security 上下文，业务接口再根据角色完成权限判断。"),
    ],
    "系统测试": [
        ("project-test-case-table.png", "图6-1 系统测试用例表", "图6-1汇总了系统测试中采用的典型用例，覆盖登录认证、心理测评、预约咨询、在线聊天、后台管理和异常访问等场景。测试用例既包含正常流程，也包含无权限访问、非法状态操作和敏感词输入等情况，可以较全面地检验系统在不同角色和不同业务状态下的表现。"),
    ],
}

CODE_BLOCKS = {
    "登录认证与权限控制": (
        "代码5-1 JWT认证过滤器关键代码",
        [
            'String authHeader = request.getHeader("Authorization");',
            'if (authHeader != null && authHeader.startsWith("Bearer ")) {',
            '    String token = authHeader.substring(7);',
            '    if (jwtTokenUtil.validateToken(token)) {',
            '        Long userId = jwtTokenUtil.getUserId(token);',
            '        String role = jwtTokenUtil.getRole(token);',
            '        UsernamePasswordAuthenticationToken authenticationToken =',
            '            new UsernamePasswordAuthenticationToken(userId, null,',
            '                List.of(new SimpleGrantedAuthority("ROLE_" + role)));',
            '        SecurityContextHolder.getContext().setAuthentication(authenticationToken);',
            '    }',
            '}',
            'filterChain.doFilter(request, response);',
        ],
        "代码5-1体现了系统身份认证的核心逻辑。过滤器会从请求头中读取 Authorization 字段，只有 token 格式正确并且通过工具类校验后，系统才会解析出用户编号和角色信息，再把认证对象放入 Spring Security 上下文。这样一来，后续业务接口不需要重复解析 token，就能直接根据上下文中的身份信息判断访问权限。"
    ),
    "心理测评与风险分析功能实现": (
        "代码5-2 心理测评提交与风险结果保存关键代码",
        [
            "AssessmentScale scale = assessmentScaleMapper.selectById(request.getScaleId());",
            'if (scale == null || scale.getEnabled() == null || scale.getEnabled() != 1) {',
            '    throw new BusinessException("量表不存在或已停用");',
            "}",
            "Map<String, Double> features = parseFeatures(request.getAnswerJson());",
            "PredictionResult prediction = predictor.predict(scale.getCode(), features);",
            "AssessmentRecord record = new AssessmentRecord();",
            "record.setUserId(userId);",
            "record.setScore(prediction.score());",
            "record.setRiskProbability(prediction.riskProbability());",
            "record.setResultLevel(prediction.resultLevel());",
            "record.setAnalysis(prediction.analysis());",
            "record.setLeadingFactorsJson(writeFactors(prediction.leadingFactors()));",
            "assessmentRecordMapper.insert(record);",
        ],
        "代码5-2把测评业务拆成了量表校验、答案解析、模型预测和记录保存几个步骤。系统先保证用户提交的量表处于启用状态，再把答案 JSON 转换为特征变量并交给预测器处理，随后把分数、风险概率、结果等级、分析说明和主要因素写入测评记录表。这种处理方式能让测评结果既可以被用户查看，也可以被管理员用于风险监测。"
    ),
    "预约咨询功能实现": (
        "代码5-3 预约确认与拒绝状态流转关键代码",
        [
            "Appointment appointment = assertCounselorOwner(userId, appointmentId);",
            'if (!"PENDING".equals(appointment.getStatus())) {',
            '    throw new BusinessException("当前预约状态不能执行同意操作");',
            "}",
            'appointment.setStatus("CONFIRMED");',
            'appointment.setReminderStatus("READY");',
            "appointmentMapper.updateById(appointment);",
            "",
            "Appointment rejected = assertCounselorOwner(userId, appointmentId);",
            'if (!"PENDING".equals(rejected.getStatus())) {',
            '    throw new BusinessException("当前预约状态不能执行拒绝操作");',
            "}",
            'rejected.setStatus("REJECTED");',
            'rejected.setReminderStatus("CLOSED");',
            "appointmentMapper.updateById(rejected);",
        ],
        "代码5-3体现了预约状态控制的关键点。咨询师处理预约前，系统会先确认该预约确实属于当前咨询师，随后检查预约是否仍处于 PENDING 状态；只有待处理的预约才能被确认或拒绝。确认后系统会把预约设置为 CONFIRMED 并开启提醒与聊天准备状态，拒绝后则关闭后续提醒，这样可以避免已经处理过的预约被重复操作。"
    ),
    "在线聊天功能实现": (
        "代码5-4 聊天准入与敏感词处理关键代码",
        [
            "public ChatMessage saveMessage(Long senderId, Long receiverId, String content, String fileUrl) {",
            "    assertChatAllowed(senderId, receiverId);",
            "    ChatMessage message = new ChatMessage();",
            "    message.setSenderId(senderId);",
            "    message.setReceiverId(receiverId);",
            "    message.setContent(maskSensitive(content));",
            "    message.setFileUrl(fileUrl);",
            "    message.setSensitiveFlag(containsSensitive(content) ? 1 : 0);",
            '    message.setReviewStatus(containsSensitive(content) ? "PENDING" : "APPROVED");',
            "    chatMessageMapper.insert(message);",
            "    return message;",
            "}",
            "private void assertChatAllowed(Long userId, Long peerId) {",
            "    if (!canChat(userId, peerId)) {",
            '        throw new BusinessException("该聊天对象尚未建立可沟通的预约关系");',
            "    }",
            "}",
        ],
        "代码5-4说明了聊天模块并不是简单地保存文本内容，而是会先判断双方是否具备沟通条件。若双方没有确认过的预约关系，也没有既有消息记录，系统会拒绝本次发送；在允许发送后，系统还会对内容进行敏感词检测与脱敏，并把待审核状态写入消息记录。该逻辑能让聊天功能和咨询预约流程保持一致，也能为管理员后续审查提供数据依据。"
    ),
}

REWRITE_FIXES = {
    "本研究": "这项研究",
    "阐述": "说明",
    "依据": "根据",
    "呈现": "显示出",
    "导致": "造成",
    "首先": "一是",
    "其次": "二是",
    "最后": "三是",
}


def is_heading(p):
    return p.style and p.style.name.startswith("Heading")


def heading_level(p):
    m = re.match(r"Heading (\d+)", p.style.name if p.style else "")
    return int(m.group(1)) if m else 0


def insert_after(paragraph, text: str | None = None, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        run = p.add_run(text)
        set_chinese_font(run)
    return p


def set_chinese_font(run, size=Pt(10.5), font="宋体"):
    run.font.name = font
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)


def set_reference_font(run, size=Pt(10.5)):
    run.font.name = "Times New Roman"
    run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")


def insert_before(paragraph, text: str | None = None, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    p = paragraph.__class__(new_p, paragraph._parent)
    if style:
        p.style = style
    if text:
        run = p.add_run(text)
        set_chinese_font(run)
    return p


def strip_numbering(p):
    ppr = p._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def find_heading(doc, text):
    for p in doc.paragraphs:
        if is_heading(p) and p.text.strip() == text:
            return p
    return None


def find_reference_heading(doc):
    for p in doc.paragraphs:
        if p.text.strip() == "参考文献":
            return p
    return None


def delete_after(paragraph):
    node = paragraph._p.getnext()
    while node is not None:
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt


def section_has_body(doc, heading):
    paras = doc.paragraphs
    idx = next(i for i, p in enumerate(paras) if p._p is heading._p)
    base = heading_level(heading)
    for p in paras[idx + 1:]:
        if is_heading(p) and heading_level(p) <= base:
            return False
        if p.text.strip():
            return True
        if p._p.xpath(".//w:drawing"):
            return True
    return False


def insert_paragraphs_after(anchor, paragraphs):
    cursor = anchor
    for text in paragraphs:
        cursor = insert_after(cursor, text, "Normal")
        cursor.paragraph_format.first_line_indent = Pt(21)
        cursor.paragraph_format.line_spacing = 1.5
    return cursor


def insert_picture_after(anchor, image_name, caption, explanation, width=5.8):
    cursor = anchor
    image_path = ASSETS / image_name
    p_img = insert_after(cursor, None, "Normal")
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(image_path), width=Inches(width))

    p_cap = insert_after(p_img, caption, "Normal")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_cap.runs:
        set_chinese_font(run, Pt(10.5), "宋体")

    p_explain = insert_after(p_cap, explanation, "Normal")
    p_explain.paragraph_format.first_line_indent = Pt(21)
    p_explain.paragraph_format.line_spacing = 1.5
    for run in p_explain.runs:
        set_chinese_font(run)
    return p_explain


def add_code_block(anchor, title, lines, explanation):
    cursor = insert_after(anchor, title, "Normal")
    cursor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cursor.runs:
        set_chinese_font(run)

    code_p = insert_after(cursor, None, "Normal")
    code_p.paragraph_format.left_indent = Pt(18)
    code_p.paragraph_format.right_indent = Pt(12)
    code_p.paragraph_format.line_spacing = 1.0
    run = code_p.add_run("\n".join(lines))
    set_chinese_font(run, Pt(8), "Consolas")

    explain_p = insert_after(code_p, explanation, "Normal")
    explain_p.paragraph_format.first_line_indent = Pt(21)
    explain_p.paragraph_format.line_spacing = 1.5
    for run in explain_p.runs:
        set_chinese_font(run)
    return explain_p


def rebuild_references(doc):
    ref_heading = find_reference_heading(doc)
    if ref_heading is None:
        ref_heading = doc.add_heading("参考文献", level=1)
    page_break = insert_before(ref_heading, None, "Normal")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    delete_after(ref_heading)
    cursor = ref_heading
    for index, item in enumerate(REFERENCES, start=1):
        cursor = insert_after(cursor, f"[{index}] {item}", "Normal")
        strip_numbering(cursor)
        cursor.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cursor.paragraph_format.first_line_indent = None
        cursor.paragraph_format.left_indent = Pt(0)
        cursor.paragraph_format.line_spacing = 1.0
        for run in cursor.runs:
            set_reference_font(run, Pt(10.5))


def normalize_text(doc):
    for p in doc.paragraphs:
        if p.text.strip().startswith("[") and re.match(r"^\[\d+\]", p.text.strip()):
            continue
        for run in p.runs:
            text = run.text
            for old, new in REWRITE_FIXES.items():
                text = text.replace(old, new)
            run.text = text


def add_content(doc):
    if find_heading(doc, "系统结构设计") is None:
        parent = find_heading(doc, "系统总体设计")
        if parent:
            h2 = insert_after(parent, "系统结构设计", "Heading 2")
            for run in h2.runs:
                set_chinese_font(run, Pt(14), "黑体")
            insert_paragraphs_after(h2, [
                "系统结构设计主要用于说明平台内部各层之间的组织关系。根据前文需求分析，平台需要同时支撑学生端、咨询师端和管理员端操作，因此系统在结构上采用前后端分离与分层处理相结合的方式，把页面展示、接口控制、业务处理、数据访问和数据存储分别放在不同层次中完成。"
            ])

    for title, intro in SECTION_INTROS.items():
        h = find_heading(doc, title)
        if h and not section_has_body(doc, h):
            insert_paragraphs_after(h, [intro])

    for title, paragraphs in EXTRA_SECTIONS.items():
        h = find_heading(doc, title)
        if h:
            insert_paragraphs_after(h, paragraphs)

    for title, items in FIGURES.items():
        h = find_heading(doc, title)
        if h:
            cursor = h
            for image_name, caption, explanation in reversed(items):
                # Insert in reverse because each insertion goes immediately after the heading.
                cursor = insert_picture_after(h, image_name, caption, explanation)

    for title, block in CODE_BLOCKS.items():
        h = find_heading(doc, title)
        if h:
            add_code_block(h, *block)


def write_report():
    lines = [
        "原文：参考文献页出现自动编号与手工编号叠加的问题，如“[3] [3]”。",
        "修改后：已重建参考文献区域，统一保留三十条文献的手工编号格式，不再使用自动编号样式。",
        "具体位置：参考文献部分。",
        "修改原因：修复编号重复问题，避免论文排版中出现截图所示的编号叠加现象。",
        "",
        "原文：用例模型分析小节正文展开不足。",
        "修改后：补充了学生、咨询师、管理员三类参与者的用例说明，并加入系统用例模型图及文字解释。",
        "具体位置：第三章，第“用例模型分析”节。",
        "修改原因：补足需求分析章节的模型说明，使功能边界和角色权限更清楚。",
        "",
        "原文：系统功能设计、系统流程设计、数据库设计等部分图文支撑不足。",
        "修改后：加入系统结构图、系统架构图、系统功能图、系统流程图、数据流图、ER图、数据库表结构设计图和接口设计表，并分别补充可写入论文的说明文字。",
        "具体位置：第四章，系统总体设计相关小节。",
        "修改原因：增强总体设计章节的图表密度和论文表达完整度。",
        "",
        "原文：核心功能实现部分主要是文字说明，缺少与项目代码对应的关键实现展示。",
        "修改后：加入 JWT 认证过滤器、测评提交与风险结果保存、预约状态流转、聊天准入与敏感词处理四处关键代码片段，并补充实现逻辑解释。",
        "具体位置：第五章，登录认证与权限控制、心理测评与风险分析功能实现、预约咨询功能实现、在线聊天功能实现等小节。",
        "修改原因：把论文内容与实际项目代码对应起来，提高实现章节的可信度和工作量表现。",
        "",
        "原文：测试方法小节正文展开不足。",
        "修改后：补充功能测试、接口测试、异常测试和兼容性测试的测试方法说明，并加入系统测试用例表。",
        "具体位置：第六章，第“测试方法”和“系统测试”相关小节。",
        "修改原因：补全测试章节，使测试过程更符合毕业论文写作要求。",
        "",
        "原文：部分句式较短，且存在“本研究、阐述、依据、呈现、导致、首先、其次、最后”等较标准化表达。",
        "修改后：已按要求替换为“这项研究、说明、根据、显示出、造成、一是、二是、三是”等表达，并适度合并短句。",
        "具体位置：全文正文段落。",
        "修改原因：按照用户给出的语言调整规则降低生硬感，同时保留论文正式表达。",
    ]
    REPORT.write_text("\n".join(lines), encoding="utf-8")


def stats(doc):
    text = "\n".join(p.text for p in doc.paragraphs)
    cjk = len(re.findall(r"[\u4e00-\u9fff]", text))
    refs = len([p for p in doc.paragraphs if re.match(r"^\[\d+\]\s", p.text.strip())])
    drawings = sum(len(p._p.xpath(".//w:drawing")) for p in doc.paragraphs)
    duplicated_ref = [p.text for p in doc.paragraphs if re.match(r"^\[\d+\]\s+\[\d+\]", p.text.strip())]
    return cjk, refs, drawings, duplicated_ref[:5]


def main():
    doc = Document(SRC)
    normalize_text(doc)
    add_content(doc)
    rebuild_references(doc)
    doc.save(OUT_TMP)
    shutil.copyfile(OUT_TMP, OUT_CN)
    write_report()
    final = Document(OUT_TMP)
    cjk, refs, drawings, dup = stats(final)
    print(f"saved={OUT_TMP.name}")
    print(f"saved_cn={OUT_CN.name}")
    print(f"cjk={cjk}")
    print(f"references={refs}")
    print(f"drawings={drawings}")
    print(f"duplicate_reference_samples={dup}")


if __name__ == "__main__":
    main()
