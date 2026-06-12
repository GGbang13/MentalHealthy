from __future__ import annotations

from pathlib import Path
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "project_current-1_ER和数据表修改版.docx"
OUTPUT_DOCX = ROOT / "project_current-1_功能结构图精简版.docx"
ASSET_DIR = ROOT / "docs" / "thesis-assets"
FUNCTION_IMAGE = ASSET_DIR / "project-system-function.png"
DOC_PARENT = None


MODULES = [
    ("账号管理", ["注册登录", "资料维护", "权限识别"]),
    ("心理测评", ["量表作答", "结果生成", "记录查看", "风险分析"]),
    ("咨询服务", ["咨询师列表", "详情查看", "评价管理"]),
    ("预约咨询", ["提交预约", "状态跟踪", "预约处理"]),
    ("在线沟通", ["联系人列表", "消息收发", "记录保存"]),
    ("文章通知", ["文章浏览", "通知查看", "内容发布"]),
    ("咨询师端", ["资料维护", "预约审核", "来访沟通"]),
    ("管理员端", ["用户管理", "内容管理", "风险监控", "数据统计"]),
]


def get_font(size: int):
    for candidate in [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
    ]:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def text_center(draw, box, text, font):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=font, fill="black", spacing=6, align="center")


def draw_box(draw, box, text, font, width=3):
    draw.rectangle(box, outline="black", width=width)
    text_center(draw, box, text, font)


def generate_function_tree():
    img = Image.new("RGB", (2600, 1020), "white")
    draw = ImageDraw.Draw(img)
    root_font = get_font(34)
    module_font = get_font(29)
    leaf_font = get_font(25)

    root = (1140, 40, 1460, 120)
    draw_box(draw, root, "心理健康服务平台", root_font, width=3)

    trunk_y = 175
    draw.line([(1300, 120), (1300, trunk_y)], fill="black", width=3)

    module_centers = [185 + i * 320 for i in range(len(MODULES))]
    left_x, right_x = module_centers[0], module_centers[-1]
    draw.line([(left_x, trunk_y), (right_x, trunk_y)], fill="black", width=3)

    module_top = 215
    module_w, module_h = 200, 75
    leaf_top = 380
    leaf_w, leaf_h = 70, 210
    leaf_gap = 18

    for center_x, (module, leaves) in zip(module_centers, MODULES):
        draw.line([(center_x, trunk_y), (center_x, module_top)], fill="black", width=3)
        mbox = (center_x - module_w // 2, module_top, center_x + module_w // 2, module_top + module_h)
        draw_box(draw, mbox, module, module_font, width=3)
        leaf_count = len(leaves)
        total_leaf_w = leaf_count * leaf_w + (leaf_count - 1) * leaf_gap
        start_x = center_x - total_leaf_w / 2
        branch_y = 345
        draw.line([(center_x, module_top + module_h), (center_x, branch_y)], fill="black", width=2)
        draw.line([(start_x + leaf_w / 2, branch_y), (start_x + total_leaf_w - leaf_w / 2, branch_y)], fill="black", width=2)
        for idx, leaf in enumerate(leaves):
            lx = start_x + idx * (leaf_w + leaf_gap) + leaf_w / 2
            draw.line([(lx, branch_y), (lx, leaf_top)], fill="black", width=2)
            label = "\n".join(list(leaf))
            lbox = (lx - leaf_w / 2, leaf_top, lx + leaf_w / 2, leaf_top + leaf_h)
            draw_box(draw, lbox, label, leaf_font, width=2)

    img.save(FUNCTION_IMAGE)


def add_run(paragraph, text, size=10.5):
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    return run


def format_body(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def insert_after(element, text="", style=None, align=None, first_line=True):
    p = OxmlElement("w:p")
    element.addnext(p)
    para = Paragraph(p, DOC_PARENT)
    if style:
        para.style = style
    if align is not None:
        para.alignment = align
    if text:
        add_run(para, text)
    format_body(para, first_line=first_line)
    return para


def rebuild_function_section():
    global DOC_PARENT
    shutil.copy2(INPUT_DOCX, OUTPUT_DOCX)
    doc = Document(str(OUTPUT_DOCX))
    DOC_PARENT = doc._body

    target = None
    next_h2 = None
    seen = False
    for p in doc.paragraphs:
        if p.text.strip() in ("系统结构与功能设计", "系统功能结构设计"):
            target = p
            seen = True
            continue
        if seen and p.style.name.startswith("Heading 2"):
            next_h2 = p
            break
    if target is None or next_h2 is None:
        raise RuntimeError("Cannot locate function section")

    target.text = "系统功能结构设计"
    target.style = "Heading 2"

    body = doc._body._element
    children = list(body)
    start = children.index(target._p)
    end = children.index(next_h2._p)
    for child in children[start + 1:end]:
        body.remove(child)

    p = insert_after(target._p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p.add_run().add_picture(str(FUNCTION_IMAGE), width=Cm(15.8))
    after = p._p
    cap = insert_after(after, "图4-1 系统功能结构图", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    after = cap._p
    text1 = (
        "如图4-1所示，系统功能结构以心理健康服务过程和角色使用场景为主线进行划分，整体包括账号管理、心理测评、咨询服务、预约咨询、在线沟通、文章通知、咨询师端和管理员端等模块。"
        "账号管理模块负责注册登录、资料维护和权限识别，是用户进入平台后使用其他功能的基础；心理测评模块负责量表作答、结果生成、记录查看和风险分析，能够把用户的测评行为转化为可保存、可追溯的结果数据。"
    )
    p = insert_after(after, text1)
    after = p._p
    text2 = (
        "咨询服务模块主要包括咨询师列表、详情查看和评价管理，预约咨询模块主要包括提交预约、状态跟踪和预约处理，在线沟通模块负责联系人列表、消息收发和记录保存。"
        "文章通知模块承担文章浏览、通知查看和内容发布任务；咨询师端面向咨询师角色，提供资料维护、预约审核和来访沟通功能；管理员端面向管理员角色，提供用户管理、内容管理、风险监控和数据统计功能。"
        "精简后的功能图保留了系统的主要功能边界，避免过多细项造成图中文字重叠，也便于在论文中清晰说明各模块之间的关系。"
    )
    insert_after(after, text2)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    generate_function_tree()
    rebuild_function_section()
    print(f"image: {FUNCTION_IMAGE}")
    print(f"docx: {OUTPUT_DOCX}")
