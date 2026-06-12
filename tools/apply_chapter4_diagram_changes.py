from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "project_current_merge_work.docx"
TMP = ROOT / "project_current_ch4_text.docx"
OUT = ROOT / "project_current-1_第四章图表修改版.docx"
FUNCTION = ROOT / "docs" / "thesis-assets" / "project-system-function.png"
FLOW = ROOT / "docs" / "thesis-assets" / "project-system-flow.png"


def set_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None


def adjust_picture_width(paragraph, cx, cy):
    for extent in paragraph._p.xpath(".//wp:extent"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))
    for extent in paragraph._p.xpath(".//a:ext"):
        extent.set("cx", str(cx))
        extent.set("cy", str(cy))


def main():
    doc = Document(SRC)

    # 修改 4-1：结构图与功能图合并。
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "系统结构设计":
            set_text(p, "系统结构与功能设计")
        elif p.text.strip() == "图4-1 系统结构图":
            set_text(p, "图4-1 系统结构与功能图")
        elif p.text.strip().startswith("如图4-1所示"):
            set_text(
                p,
                "如图4-1所示，系统结构与功能设计被合并到同一张图中进行说明。图的上方说明平台采用表现层、业务服务层、数据访问层、数据存储层和模型分析层等结构层次；图的下方按照用户、心理咨询师和管理员三类角色展开功能模块。用户侧主要包括账号管理、心理测评、咨询预约和互动服务；心理咨询师侧主要包括资料维护、预约处理和咨询沟通；管理员侧主要包括用户管理、内容管理和数据监管。该图既能说明系统分层，也能说明各角色在平台中的主要操作范围。"
            )
        elif p.text.strip().startswith("系统结构设计主要用于说明"):
            set_text(
                p,
                "合并后的图能够避免结构图和功能图重复表达的问题，也能更直观地表现前端页面、后端业务、数据库和风险预测模型之间的配合关系。前端按照页面和用户角色组织组件，后端按照业务模块划分 Controller、Service、Mapper 和实体对象，数据库表则对应到具体业务对象，从而使系统结构和功能边界更加清楚。"
            )

    # 放大 4-1 合并图和 4-4 业务流程图。
    for idx, p in enumerate(doc.paragraphs):
        if idx == 191 and p._p.xpath(".//w:drawing"):
            adjust_picture_width(p, 5760720, 3323520)
        if idx == 212 and p._p.xpath(".//w:drawing"):
            adjust_picture_width(p, 5760720, 4320540)

    # 修改流程图题注与说明。
    for p in doc.paragraphs:
        if p.text.strip() == "心理健康服务平台系统流程图":
            set_text(p, "图4-4 系统业务流程图")
        elif p.text.strip().startswith("图4.3说明了用户进入系统后的主要业务流程"):
            set_text(
                p,
                "图4-4展示了平台的系统业务流程。用户进入系统后首先输入账号和密码，系统判断登录是否成功；若登录失败，则提示错误并返回登录；若登录成功，则继续判断用户角色。普通用户进入用户端首页后，可以选择是否进行心理测评，若选择测评则提交量表答案并生成风险等级与测评记录，随后再判断是否预约心理咨询；若选择预约，则提交预约申请并等待咨询师处理。咨询师进入工作台后查看预约申请，并判断是否同意预约，同意后开放在线聊天，拒绝后关闭该预约。管理员进入后台后查看测评与预约数据，并判断是否存在异常记录，若存在异常则进行审核处理，若不存在则继续维护用户、文章和数据面板。该流程图包含登录判断、角色判断、测评判断、预约判断、咨询师审核判断和管理员异常判断，因此能够体现系统业务流转中的关键分支。"
            )

    # 删除原 4-3 系统功能图小节：Heading、图片、题注、两段说明。
    remove_texts = {
        "系统功能设计",
        "系统功能图",
    }
    to_delete = []
    for i, p in enumerate(doc.paragraphs):
        if 205 <= i <= 210:
            to_delete.append(p)
    for p in to_delete:
        delete_paragraph(p)

    doc.save(TMP)

    # 替换图片二进制：image6 为合并图，image9 为业务流程图。
    with ZipFile(TMP, "r") as zin, ZipFile(OUT, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/media/image6.png":
                data = FUNCTION.read_bytes()
            elif item.filename == "word/media/image9.png":
                data = FLOW.read_bytes()
            zout.writestr(item, data)

    print(OUT)


if __name__ == "__main__":
    main()
