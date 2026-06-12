from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("project_current_2_cjk20000_review.docx")
OUT = Path("project_current_2_中文字符约20000版.docx")
REPORT = Path("docs/project_current_2_reference_locations_20000cjk.txt")


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = style
    except Exception:
        new_para.style = "Normal"
    new_para.add_run(text)
    return new_para


def find_para(doc: Document, exact: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact:
            return para
    raise ValueError(exact)


def main() -> None:
    doc = Document(SRC)
    additions = {
        "前端工程设计": [
            "前端工程设计还需要关注页面复用和状态一致性。系统中登录状态、用户角色、用户资料等信息会被多个页面使用，如果每个页面都单独请求和保存，容易造成状态不同步。因此，系统使用 Pinia 统一保存 token 和用户信息，并在 Axios 请求拦截器中统一注入认证头。页面组件只需要关注自身业务数据，不需要重复处理登录态细节。这样的设计减少了页面之间的耦合，也使后续新增页面时能够复用同一套状态管理逻辑[28-29]。",
            "在页面组织上，系统将登录页、布局页、用户端页面、咨询师端页面和管理员端页面分开实现。布局页负责导航、侧边栏和主体区域，业务页面负责具体功能。对于管理员文章新增和编辑这类相似页面，可以通过表单组件复用减少重复代码；对于测评、预约和聊天这类交互较强的页面，则将接口调用、表单校验和结果展示集中在对应视图中处理。这样的工程组织方式有助于保持前端项目结构清晰[13-14][28]。",
        ],
        "后端接口与数据交互设计": [
            "后端接口设计遵循按业务模块划分的原则。认证相关接口统一放在 `/api/auth` 下，用户资料接口放在 `/api/users` 下，咨询师接口放在 `/api/counselors` 下，预约接口放在 `/api/appointments` 下，测评接口放在 `/api/assessments` 下，聊天接口放在 `/api/chat` 下，后台管理接口放在 `/api/admin` 下。这样的路径划分使接口语义更加直观，也便于前端按照模块封装请求方法[18][29]。",
            "数据交互方面，前端通过统一 HTTP 模块调用后端接口，后端返回统一响应结构。当前端收到业务成功响应时，直接取出 data 字段渲染页面；当收到未授权响应时，前端清空用户状态并跳转登录页；当收到业务异常时，页面通过消息提示展示错误原因。统一响应和统一错误处理能够减少页面中的重复判断，也便于后续统一调整错误提示和登录失效处理逻辑[15-18]。",
        ],
        "测试结果分析": [
            "从功能覆盖角度看，本系统测试重点围绕三类角色和六条主流程展开。三类角色分别是普通用户、心理咨询师和管理员；六条主流程分别是登录认证流程、心理测评流程、预约咨询流程、在线聊天流程、内容管理流程和数据看板流程。测试时不仅需要验证每个按钮是否可点击，还要验证前后端数据是否一致，例如预约确认后用户端是否出现聊天入口，测评提交后管理员端是否可以查看记录，文章发布后用户端是否可以阅读[13-14]。",
            "从异常处理角度看，测试需要关注权限越界、参数缺失、数据不存在和网络异常等情况。权限越界主要包括普通用户访问管理员接口、咨询师处理不属于自己的预约、用户访问他人聊天记录等；参数缺失主要包括预约时间为空、测评答案为空、文章标题为空等；数据不存在主要包括访问已删除记录或无效编号；网络异常主要包括接口超时和 WebSocket 断开。通过这些测试，可以发现系统在边界场景下是否仍能给出明确反馈[25-26]。",
        ],
    }
    notes = []
    for heading, paras in additions.items():
        current = find_para(doc, heading)
        for text in paras:
            current = insert_after(current, text)
        notes.append(f"{heading}：第二轮补充 {len(paras)} 段，用于使中文字符数接近 20000。")
    doc.save(OUT)
    with REPORT.open("a", encoding="utf-8") as f:
        f.write("\n\n第二轮补写位置：\n" + "\n".join(notes))
    print(OUT)


if __name__ == "__main__":
    main()
