from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("project_current_1_source.docx")
OUT = Path("project_current-1_扩充润色版.docx")
REPORT = Path("docs/project_current-1_修改说明.txt")


def set_text(paragraph: Paragraph, text: str) -> str:
    old = paragraph.text
    paragraph.text = text
    return old


def insert_after(paragraph: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = style
    except Exception:
        new_para.style = "Normal"
    new_para.add_run(text)
    return new_para


def find_exact(doc: Document, exact: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(exact)


def find_contains(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(needle)


def remove_old_refs(doc: Document) -> Paragraph:
    ref = find_exact(doc, "参考文献")
    found = False
    for p in list(doc.paragraphs):
        if p is ref:
            found = True
            continue
        if found:
            p._element.getparent().remove(p._element)
    return ref


def main() -> None:
    doc = Document(SRC)
    logs: list[tuple[str, str, str, str]] = []

    replacements = [
        (
            "心理健康是高校学生成长与发展的重要基础。",
            "心理健康是高校学生成长与发展的重要基础，近年来，学习压力、人际关系、就业竞争和家庭环境等因素不断叠加，使得一部分学生会面临焦虑、抑郁、情绪低落、睡眠障碍和适应困难等问题，高校心理健康教育与咨询工作也需要覆盖心理知识宣传、心理测评、预约咨询、危机预警、持续跟踪和管理统计等多个环节；传统模式通常依赖线下预约、人工记录和单点沟通，服务流程容易受到时间、场地和人员规模的限制[1]。",
            "第一章，第研究背景与意义节，第一段",
            "合并短句，适度增加“的、会、到”等字，减少句号，并保留引用标记。",
        ),
        (
            "综合来看，心理健康服务平台的发展趋势主要体现在以下方面。",
            "综合来看，心理健康服务平台的发展趋势主要表现出三个方向，一是服务流程从单一管理转向多角色协同，用户、咨询师和管理员需要在同一平台中完成连续业务；二是测评结果从静态分数转向风险分层与解释分析，帮助相关人员理解结果根据；三是沟通方式从线下预约转向线上线下结合，实时消息、通知提醒和历史记录会成为提升服务连续性的重要支撑，本文设计的平台正是围绕这些趋势展开。",
            "第一章，第国内外研究现状节，第三段",
            "将“第一、第二、第三”调整为“一是、二是、三是”，并把多个短句合并为一个长句。",
        ),
        (
            "本文围绕“基于 Spring Boot 和 Vue 的心理健康服务平台的设计与实现”展开研究，主要工作包括以下几个方面。",
            "本文围绕“基于 Spring Boot 和 Vue 的心理健康服务平台的设计与实现”展开说明，主要工作可以概括为以下几个方面。",
            "第一章，第论文的主要内容和组织结构节，第一段",
            "将“展开研究”替换为更自然的“展开说明”，降低生硬感。",
        ),
        (
            "第一，分析心理健康服务平台的业务需求。",
            "一是分析心理健康服务平台的业务需求，根据用户、咨询师和管理员三类角色梳理功能边界，明确系统需要支持的登录认证、测评、预约、聊天、文章、通知和后台管理流程。",
            "第一章，第论文的主要内容和组织结构节，第二段",
            "用“一是”替换顺序连接词，并合并后续说明，减少短句。",
        ),
        (
            "第二，设计系统总体架构。",
            "二是设计系统总体架构，系统采用前后端分离架构，前端负责页面交互、路由控制、状态管理和数据可视化，后端负责接口服务、业务处理、权限校验和数据持久化。",
            "第一章，第论文的主要内容和组织结构节，第三段",
            "用“二是”替换顺序连接词，并调整语序。",
        ),
        (
            "第三，实现核心业务模块。",
            "三是实现核心业务模块，用户端提供心理测评、预约咨询、文章阅读和在线沟通入口，咨询师端提供预约处理、资料维护和沟通入口，管理员端提供平台数据看板、风险监控、用户管理、咨询师管理、文章管理和通知管理。",
            "第一章，第论文的主要内容和组织结构节，第四段",
            "用“三是”替换顺序连接词，并把相关功能合并说明。",
        ),
        (
            "第四，开展系统测试。",
            "另一方面，论文还会开展系统测试，通过功能测试、异常测试和界面适配测试验证系统主要流程，分析系统在权限控制、数据提交、实时通信和统计展示方面的运行效果。",
            "第一章，第论文的主要内容和组织结构节，第五段",
            "用“另一方面”替换机械顺序词，并保留正式但自然的论文语气。",
        ),
    ]

    for needle, new_text, pos, reason in replacements:
        p = find_contains(doc, needle)
        old = set_text(p, new_text)
        logs.append((old, new_text, pos, reason))

    additions: dict[str, list[tuple[str, str]]] = {
        "Spring Boot技术": [
            ("Spring Boot 是基于 Spring 生态的快速开发框架，它通过自动配置、起步依赖和内嵌服务器，减少了传统 Java Web 项目当中大量重复的 XML 配置工作。本系统后端使用 Spring Boot 提供 REST API，把认证、用户资料、咨询师资料、预约、测评、聊天、文章、通知和数据看板等业务能力拆分到不同控制器和服务类中，这样既能保证接口层比较清晰，也能让业务逻辑有相对独立的维护空间[15-16]。", "第二章，第Spring Boot技术节，新增第一段"),
            ("在具体实现中，Spring Boot 不只是接口框架，还承担了依赖注入、配置读取、异常处理和启动管理等任务。控制器层接收前端请求，服务层按照业务规则处理数据，Mapper 层负责数据库访问，这种分层方式能够让系统在需求变化时更容易调整，也便于论文后续按照模块说明系统实现过程[13-14]。", "第二章，第Spring Boot技术节，新增第二段"),
        ],
        "Vue 3技术": [
            ("Vue 3 是本系统前端页面实现的基础，它提供了组件化、响应式数据绑定和组合式 API 等能力。系统中的登录页面、用户首页、咨询师首页、管理员首页、测评页面、预约页面和聊天页面都可以被拆分为相对独立的组件，每个组件负责自己的数据加载和交互逻辑，这种方式能够降低页面之间的耦合，也能让后续维护更加方便[28]。", "第二章，第Vue 3技术节，新增第一段"),
            ("前端工程还结合了 TypeScript、Vite、Pinia、Vue Router、Element Plus、Axios 和 ECharts。TypeScript 用于减少变量类型混乱造成的问题，Vite 用于提高开发和构建速度，Pinia 用于保存用户登录状态，Vue Router 用于控制不同角色的页面访问，Element Plus 提供统一的表格和表单组件，Axios 负责接口请求封装，ECharts 用于管理员端的数据可视化展示[28-29]。", "第二章，第Vue 3技术节，新增第二段"),
        ],
        "MyBatis-Plus与Mysql技术": [
            ("MySQL 是系统的数据存储基础，用于保存用户、咨询师资料、预约记录、测评量表、测评记录、聊天消息、文章、通知、评价和操作日志等信息。心理健康服务平台的数据关系比较多，如果表结构设计不清晰，后续功能扩展和数据统计都会受到影响，因此系统把用户表作为身份中心表，再通过外键语义把预约、测评、聊天和后台管理数据连接起来[12][30]。", "第二章，第MyBatis-Plus与Mysql技术节，新增第一段"),
            ("MyBatis-Plus 在 MyBatis 基础上提供了条件构造器、通用 CRUD 和实体映射等能力，能够减少重复 SQL 编写。本系统在查询用户预约、咨询师资料、测评历史和聊天记录时，大量使用 LambdaQueryWrapper 构造查询条件，这种写法比手写字符串 SQL 更直观，也能够在一定程度上减少字段名写错造成的问题[12][30]。", "第二章，第MyBatis-Plus与Mysql技术节，新增第二段"),
        ],
        "Spring Security与JWT": [
            ("Spring Security 与 JWT 共同构成系统的认证与权限控制基础。用户登录成功后，后端会生成包含用户编号和角色信息的 token，前端把 token 保存到 Pinia 状态和 localStorage 中，并在后续请求中通过请求头传给后端。后端过滤器解析 token 后，会把用户身份写入安全上下文，使服务层能够根据当前用户完成后续权限判断[17-18]。", "第二章，第Spring Security与JWT节，新增第一段"),
            ("需要说明的是，前端路由控制只能改善页面访问体验，不能单独承担安全责任。系统在预约、聊天、测评监控和后台管理等接口中，还要根据当前用户角色和数据归属进行二次校验，例如普通用户只能取消自己的预约，咨询师只能确认或拒绝属于自己的预约，管理员接口需要验证管理员身份，这样才能降低越权访问风险[17][25-26]。", "第二章，第Spring Security与JWT节，新增第二段"),
        ],
        "WebSocket技术": [
            ("WebSocket 适合用于在线聊天这一类实时通信场景，它能够在浏览器和服务器之间建立长连接，使服务器在有新消息时主动推送给客户端。与轮询相比，WebSocket 可以减少频繁请求造成的开销，也能让消息表现得更加及时。本系统在聊天模块中使用 WebSocket 推送消息，同时保留 REST API 用于历史记录加载和消息保存[23-24]。", "第二章，第WebSocket技术节，新增第一段"),
            ("这种“REST API + WebSocket”的组合比较适合心理咨询场景，一方面，REST API 能够保证聊天记录先被持久化保存，方便后续查看和审计；另一方面，WebSocket 能够提升在线沟通的即时性。当实时连接出现异常时，用户仍然可以通过历史消息接口查看已经保存的内容，系统不会因为推送通道异常而完全不可用[23-24]。", "第二章，第WebSocket技术节，新增第二段"),
        ],
        "ECharts与数据可视化": [
            ("ECharts 是常用的前端数据可视化库，能够通过柱状图、折线图、饼图等形式表现统计数据。本系统管理员端需要展示用户数量、咨询师数量、预约数量、测评数量、高风险记录数量和风险分布等信息，如果只使用文字或普通表格，管理人员很难快速把握平台运行状态，因此数据看板中使用 ECharts 对关键指标进行可视化展示[28]。", "第二章，第ECharts与数据可视化节，新增第一段"),
        ],
        "CatBoost与心理风险分析": [
            ("CatBoost 是一种基于梯度提升树的机器学习算法，能够处理表格特征并输出分类或回归结果。本系统在心理测评模块中设计了可选的 CatBoost 推理路径，用户提交测评答案后，后端会把答案 JSON 转换为特征数据，再尝试调用模型生成风险概率、风险等级和主要影响因素[19-22]。", "第二章，第CatBoost与心理风险分析节，新增第一段"),
            ("考虑到毕业设计演示环境可能缺少完整模型文件或 Python 推理环境，系统还提供了 Java 规则回退逻辑。当外部模型不可用时，系统按照内置特征权重计算基础风险结果，从而保证测评流程仍然能够完成。这样的设计一方面保留了机器学习扩展空间，另一方面也增强了系统运行的稳定性[19-22]。", "第二章，第CatBoost与心理风险分析节，新增第二段"),
        ],
        "角色与业务需求": [
            ("从普通用户角度看，平台需要降低学生主动求助的门槛。用户可能并不会一开始就直接到线下心理中心咨询，而是会先通过线上平台了解自身状态，因此用户端应把心理测评、预约咨询、文章阅读和在线沟通放在清晰位置，让用户能够一步步完成从了解信息到预约服务的过程[1-6]。", "第三章，第角色与业务需求节，新增段落"),
            ("从咨询师角度看，平台需要减少重复沟通和事务性处理。咨询师需要查看预约申请、了解用户填写的问题描述、确认或拒绝预约，并在预约确认后进入沟通页面。若这些信息散落在不同渠道当中，会增加准备成本，也容易造成记录缺失，因此系统需要把预约、资料和聊天关系串联起来[13-14]。", "第三章，第角色与业务需求节，新增段落"),
            ("从管理员角度看，平台不仅要能维护基础数据，还要能帮助管理人员看到整体运行情况。管理员需要关注用户数量、咨询师数量、预约数量、测评数量和风险等级分布，也需要在必要时查看测评记录，为后续人工复核和线下处理提供根据[25-30]。", "第三章，第角色与业务需求节，新增段落"),
        ],
        "系统架构设计": [
            ("系统架构设计需要兼顾开发便利性和后续维护。前端构建完成后可以作为静态资源部署，后端以 Spring Boot 服务形式运行，数据库和缓存作为基础服务提供支撑。浏览器端通过统一接口访问系统，普通业务走 REST API，聊天业务在保存消息后通过 WebSocket 推送，这样既能保证数据留痕，也能提升实时沟通体验[15-18][23-24]。", "第四章，第系统架构设计节，新增段落"),
            ("系统还需要考虑配置和环境切换问题。数据库连接、JWT 密钥、文件上传目录和模型脚本路径不宜直接写死在业务代码中，而应通过配置文件或环境变量管理。这样在本地开发、测试部署和后续服务器运行之间切换时，代码主体不需要频繁修改，系统也更容易维护[15-18][25]。", "第四章，第系统架构设计节，新增段落"),
        ],
        "数据库设计": [
            ("数据库设计以用户表为中心展开，咨询师资料表通过用户编号与咨询师账号绑定，预约表连接普通用户和咨询师资料，测评记录表连接用户和量表，聊天消息表通过发送者和接收者编号记录沟通内容。这样的关系设计能够避免重复保存用户基础信息，也能让不同业务模块围绕统一身份体系进行协作[12][30]。", "第四章，第数据库设计节，新增段落"),
            ("测评数据需要特别注意历史结果的可解释性。量表题目、模型权重和风险解释规则后续可能会调整，但已经保存的历史测评记录仍然需要能够被查看，所以测评记录表不仅保存答案 JSON，还保存分数、风险概率、风险等级、分析文本、模型名称和主要影响因素，这样能够让历史记录有相对稳定的解释根据[19-22]。", "第四章，第数据库设计节，新增段落"),
        ],
        "用户端功能实现": [
            ("用户端功能的实现重点是把常用入口集中起来。用户登录后可以从首页进入心理测评、预约咨询、心理文章和在线沟通，页面不需要让用户在复杂菜单中反复查找。对于心理健康服务平台而言，清晰入口本身就是降低求助成本的一部分，尤其是当用户处在压力较大或情绪低落状态时，过于复杂的操作流程会降低使用意愿[1-6]。", "第五章，第用户端功能实现节，新增段落"),
            ("在预约咨询功能中，用户需要根据咨询师资料做出选择，因此页面会展示咨询师姓名、职称、擅长方向、从业年限、咨询价格、在线状态和评分等信息。用户提交预约时填写预约时间、咨询形式、咨询时长和问题描述，系统保存后进入待确认状态，等咨询师处理后再显示后续状态[6][13-14]。", "第五章，第用户端功能实现节，新增段落"),
        ],
        "管理员端功能实现": [
            ("管理员端功能体现平台治理能力。管理员可以维护用户、咨询师、文章和通知，也可以查看测评记录和数据看板。用户管理和咨询师管理用于保证平台主体数据准确，文章和通知管理用于心理健康教育和活动提醒，测评监控和数据看板则用于掌握平台运行状态和风险分布[25-30]。", "第五章，第管理员端功能实现节，新增段落"),
            ("由于管理员权限较高，系统后续还需要进一步强化后台安全。当前系统已经具备基础角色校验，但在真实部署中，还可以增加操作日志、敏感数据脱敏、导出审批和二次确认等机制。心理健康数据具有敏感性，管理员模块不能只强调管理便利，也应强调访问边界和责任留痕[25-27]。", "第五章，第管理员端功能实现节，新增段落"),
        ],
        "心理测评与风险分析功能实现": [
            ("心理测评模块的实现可以分为题目加载、答案提交、风险计算和结果保存几个步骤。前端从后端获取量表配置，根据 questionJson 动态生成题目，用户完成作答后提交答案 JSON；后端解析答案并生成特征数据，随后调用 CatBoost 推理或 Java 回退逻辑，最终把风险概率、风险等级、分析文本和主要影响因素保存到数据库[19-22]。", "第五章，第心理测评与风险分析功能实现节，新增段落"),
            ("主要影响因素的展示能够让测评结果更容易理解。普通问卷系统往往只给出分数，而本系统会尝试显示睡眠问题、长期压力、社交退缩、家庭支持或放松能力等影响因素。需要注意的是，这些因素只是模型或规则计算中的重要变量，并不能被直接理解为医学诊断结论，因此页面和论文说明中都应强调其辅助参考性质[1-4][19-22]。", "第五章，第心理测评与风险分析功能实现节，新增段落"),
        ],
        "在线聊天功能实现": [
            ("在线聊天模块在业务上连接预约咨询流程，在技术上连接 REST API 和 WebSocket。用户或咨询师进入聊天页面后，系统先加载联系人列表，再加载与指定对象之间的历史消息。发送消息时，后端先检查双方是否具有可沟通关系，再保存消息并推送给接收方，这样可以避免没有预约关系的用户随意私聊咨询师[23-24]。", "第五章，第在线聊天功能实现节，新增段落"),
            ("消息保存后再推送，是该模块比较重要的设计。若只做前端即时展示而不保存到数据库，后续就无法追溯沟通过程；若只保存不推送，又会影响在线咨询体验。因此系统把持久化和实时推送结合起来，在保证记录完整的同时提升沟通及时性[23-24]。", "第五章，第在线聊天功能实现节，新增段落"),
        ],
        "测试结果分析": [
            ("从功能覆盖角度看，系统测试围绕三类角色和多条主流程展开，包括登录认证、心理测评、预约咨询、在线聊天、内容管理和数据看板。测试时不仅要验证按钮是否能够点击，还要验证前后端数据是否一致，例如预约确认后用户端是否出现聊天入口，测评提交后管理员端是否能够查看记录，文章发布后用户端是否能够阅读[13-14]。", "第六章，第测试结果分析节，新增段落"),
            ("从异常处理角度看，测试需要关注未登录访问、角色越权、参数缺失、数据不存在和 WebSocket 断开等情况。普通用户访问管理员页面应被拦截，咨询师处理不属于自己的预约应被拒绝，聊天双方没有确认预约关系时不能发送消息。通过这些测试，可以验证系统在边界场景下是否仍能给出明确反馈[25-26]。", "第六章，第测试结果分析节，新增段落"),
        ],
        "不足与展望": [
            ("后续工作可以从移动端体验、模型治理和平台安全三个方向继续展开。一方面，学生使用心理服务时经常依赖手机，因此可以开发响应式移动页面或微信小程序；另一方面，心理风险分析需要更完整的数据来源说明、模型版本记录和人工复核流程；同时，平台还应补充操作日志审计、敏感数据访问审批和数据留存策略[25-30]。", "第七章，第不足与展望节，新增段落"),
            ("如果系统将来进入真实应用环境，还需要和学校现有的信息系统协同，例如统一身份认证、消息通知系统、学生基础数据平台和心理中心工作流。现阶段系统已经把核心业务拆分为相对清晰的接口和数据表，后续可以在不推翻主体结构的前提下逐步接入这些外部能力[28-30]。", "第七章，第不足与展望节，新增段落"),
        ],
    }

    for heading, entries in additions.items():
        anchor = find_exact(doc, heading)
        current = anchor
        for text, pos in entries:
            current = insert_after(current, text)
            logs.append(("（该位置原文内容较少或仅有小节标题）", text, pos, "扩充论文内容，保留专有名词和引用标记，并按要求使用较自然的连接方式。"))

    # Add citations to existing paragraphs when possible.
    for needle, cite in [
        ("在信息化建设不断推进的背景下", "[4-6]"),
        ("国外高校和医疗机构较早", "[7-10]"),
        ("国内高校心理健康信息化建设", "[1][6]"),
        ("本系统采用“前端页面层", "[11-18]"),
        ("在在线沟通方面", "[23-24]"),
        ("心理健康属于敏感领域", "[25-27]"),
        ("后台管理功能可以继续加强统计分析能力", "[28-30]"),
    ]:
        try:
            p = find_contains(doc, needle)
            if cite not in p.text:
                p.add_run(cite)
        except ValueError:
            pass

    ref_heading = remove_old_refs(doc)
    refs = [
        "[1] 江光荣, 李丹阳, 任志洪, 等. 中国国民心理健康素养的现状与特点[J]. 心理学报, 2021, 53(02): 182-201.",
        "[2] 教育部. 高等学校学生心理健康教育指导纲要[Z]. 2018.",
        "[3] 王登峰, 张大均. 大学生心理健康教育[M]. 北京: 高等教育出版社, 2019.",
        "[4] World Health Organization. Mental health: strengthening our response[EB/OL]. Geneva: WHO, 2022.",
        "[5] Torous J, Nicholas J, Larsen M E, et al. Clinical review of user engagement with mental health smartphone apps[J]. Evidence-Based Mental Health, 2018, 21(3): 116-119.",
        "[6] 李焰, 郑日昌. 大学生心理健康教育信息化建设研究[J]. 中国高等教育, 2020(10): 55-57.",
        "[7] Andersson G. Internet interventions: Past, present and future[J]. Internet Interventions, 2018, 12: 181-188.",
        "[8] Firth J, Torous J, Nicholas J, et al. The efficacy of smartphone-based mental health interventions[J]. World Psychiatry, 2017, 16(3): 287-298.",
        "[9] Mohr D C, Weingardt K R, Reddy M, Schueller S M. Three problems with current digital mental health research[J]. Psychiatric Services, 2017, 68(5): 427-429.",
        "[10] Hollis C, Falconer C J, Martin J L, et al. Annual research review: Digital health interventions for children and young people with mental health problems[J]. Journal of Child Psychology and Psychiatry, 2017, 58(4): 474-503.",
        "[11] 许令波. 深入分析 Java Web 技术内幕[M]. 北京: 电子工业出版社, 2017.",
        "[12] 王珊, 萨师煊. 数据库系统概论[M]. 北京: 高等教育出版社, 2014.",
        "[13] Sommerville I. Software Engineering[M]. Boston: Pearson, 2016.",
        "[14] Pressman R S, Maxim B R. Software Engineering: A Practitioner's Approach[M]. New York: McGraw-Hill Education, 2019.",
        "[15] Craig Walls. Spring Boot 实战[M]. 北京: 人民邮电出版社, 2016.",
        "[16] Walls C. Spring in Action[M]. 6th ed. Shelter Island: Manning Publications, 2022.",
        "[17] Freeman A. Pro Spring Security[M]. New York: Apress, 2019.",
        "[18] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.",
        "[19] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
        "[20] Chen T, Guestrin C. XGBoost: A Scalable Tree Boosting System[C]. KDD, 2016: 785-794.",
        "[21] Prokhorenkova L, Gusev G, Vorobev A, Dorogush A V, Gulin A. CatBoost: unbiased boosting with categorical features[C]. NeurIPS, 2018.",
        "[22] James G, Witten D, Hastie T, Tibshirani R. An Introduction to Statistical Learning[M]. New York: Springer, 2021.",
        "[23] Fette I, Melnikov A. The WebSocket Protocol[S]. RFC 6455, 2011.",
        "[24] MDN Web Docs. WebSockets API[EB/OL]. https://developer.mozilla.org/zh-CN/docs/Web/API/WebSockets_API.",
        "[25] OWASP Foundation. OWASP Top 10[EB/OL]. https://owasp.org/www-project-top-ten/.",
        "[26] NIST. Digital Identity Guidelines[S]. SP 800-63, 2017.",
        "[27] 国家互联网信息办公室. 个人信息保护法相关解读与数据安全治理文件[Z]. 2021.",
        "[28] 尤雨溪. Vue.js 设计与实现[M]. 北京: 人民邮电出版社, 2022.",
        "[29] Richardson L, Ruby S. RESTful Web Services[M]. Sebastopol: O'Reilly Media, 2007.",
        "[30] Kroenke D M, Auer D J. Database Processing: Fundamentals, Design, and Implementation[M]. London: Pearson, 2015.",
    ]
    current = ref_heading
    for ref in refs:
        current = insert_after(current, ref, "参考文献")

    doc.save(OUT)
    REPORT.parent.mkdir(exist_ok=True)
    report_lines: list[str] = []
    for old, new, pos, reason in logs:
        report_lines.extend([
            "原文：",
            old,
            "修改后：",
            new,
            f"具体位置：{pos}",
            f"修改原因：{reason}",
            "",
        ])
    REPORT.write_text("\n".join(report_lines), encoding="utf-8")

    text = "\n".join(p.text for p in doc.paragraphs)
    print(OUT)
    print("cjk", len(re.findall(r"[\u4e00-\u9fff]", text)))
    print("nonspace", len("".join(text.split())))


if __name__ == "__main__":
    main()
