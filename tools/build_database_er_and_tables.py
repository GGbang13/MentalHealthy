from __future__ import annotations

from pathlib import Path
import shutil
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "docs" / "thesis-assets"
INPUT_DOCX = ROOT / "project_current-1_第四章图表修改版.docx"
WORK_DOCX = ROOT / "project_current_db_work.docx"
OUTPUT_DOCX = ROOT / "project_current-1_ER和数据表修改版.docx"
ER_IMAGE = ASSET_DIR / "project-er-diagram-detailed.png"
DOC_PARENT = None


TABLES = [
    {
        "caption": "表4-1 用户表结构",
        "name": "sys_user",
        "cn": "用户表",
        "desc": "用户表是平台身份数据的核心表，用来保存普通用户、心理咨询师和管理员的账号、角色、联系方式、个人资料和状态信息。系统登录认证、角色权限判断、预约发起、测评记录归属、消息收发和后台操作追踪都会关联到该表，因此该表在数据库结构中处于中心位置。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("username", "varchar", "64", "否", "用户名，唯一"),
            ("password", "varchar", "255", "否", "加密后的登录密码"),
            ("email", "varchar", "128", "是", "邮箱"),
            ("phone", "varchar", "32", "是", "手机号"),
            ("role", "varchar", "32", "否", "用户角色"),
            ("nickname", "varchar", "64", "是", "昵称"),
            ("avatar", "varchar", "255", "是", "头像地址"),
            ("gender", "varchar", "16", "是", "性别"),
            ("age", "int", "", "是", "年龄"),
            ("profile", "text", "", "是", "个人简介"),
            ("status", "tinyint", "", "是", "账号状态"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-2 咨询师资料表结构",
        "name": "counselor_profile",
        "cn": "咨询师资料表",
        "desc": "咨询师资料表是用户表的扩展表，主要保存咨询师的职称、擅长方向、从业年限、个人介绍、收费标准、在线状态和可预约时间等信息。该表通过 user_id 与用户账号绑定，使系统能够在保持统一登录体系的同时，对咨询师业务资料进行单独维护。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("user_id", "bigint", "", "否", "关联用户编号"),
            ("title", "varchar", "64", "是", "咨询师职称"),
            ("specialties", "varchar", "255", "是", "擅长方向"),
            ("years_of_experience", "int", "", "是", "从业年限"),
            ("introduction", "text", "", "是", "个人介绍"),
            ("price_per_hour", "decimal", "10,2", "是", "小时咨询价格"),
            ("online_status", "tinyint", "", "是", "在线状态"),
            ("schedule_json", "text", "", "是", "排班时间配置"),
            ("rating", "decimal", "3,2", "是", "综合评分"),
            ("review_count", "int", "", "是", "评价数量"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-3 咨询评价表结构",
        "name": "counselor_review",
        "cn": "咨询评价表",
        "desc": "咨询评价表记录用户对心理咨询师服务的评分和文字评价。该表一方面可以为咨询师展示页提供评分依据，另一方面也可以辅助管理员了解服务质量，后续若需要做评价审核或服务质量统计，也可以在该表基础上扩展。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("counselor_id", "bigint", "", "否", "咨询师资料编号"),
            ("user_id", "bigint", "", "否", "评价用户编号"),
            ("rating", "int", "", "否", "评分"),
            ("content", "text", "", "是", "评价内容"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-4 预约记录表结构",
        "name": "appointment",
        "cn": "预约记录表",
        "desc": "预约记录表用于保存普通用户与心理咨询师之间的预约申请、预约时间、咨询方式、问题描述、处理状态和提醒状态。预约业务涉及申请、确认、取消、完成等状态变化，因此该表既保存业务内容，也承担流程追踪作用。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("user_id", "bigint", "", "否", "预约用户编号"),
            ("counselor_id", "bigint", "", "否", "咨询师资料编号"),
            ("appointment_time", "datetime", "", "否", "预约时间"),
            ("duration_minutes", "int", "", "是", "咨询时长"),
            ("type", "varchar", "32", "是", "咨询方式"),
            ("issue_description", "text", "", "是", "问题描述"),
            ("status", "varchar", "32", "是", "预约状态"),
            ("reminder_status", "varchar", "32", "是", "提醒状态"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-5 测评量表表结构",
        "name": "assessment_scale",
        "cn": "测评量表表",
        "desc": "测评量表表保存 PHQ-9、GAD-7 等心理测评量表的基本信息、题目配置和规则配置。题目与规则采用 JSON 字段保存，能够支持前端动态渲染题目，也便于后续增加新的量表类型，而不必频繁修改固定字段结构。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("name", "varchar", "128", "否", "量表名称"),
            ("code", "varchar", "64", "否", "量表编码，唯一"),
            ("description", "text", "", "是", "量表说明"),
            ("question_json", "json", "", "是", "题目配置"),
            ("rule_json", "json", "", "是", "结果规则配置"),
            ("enabled", "tinyint", "", "是", "是否启用"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-6 测评记录表结构",
        "name": "assessment_record",
        "cn": "测评记录表",
        "desc": "测评记录表保存用户每一次测评提交后的答案、得分、风险概率、结果等级、分析文本、模型名称和主要影响因素。该表既是用户查看历史测评结果的数据来源，也是管理员进行风险监控和统计分析的重要依据。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("user_id", "bigint", "", "否", "测评用户编号"),
            ("scale_id", "bigint", "", "否", "量表编号"),
            ("answer_json", "json", "", "是", "用户答案"),
            ("score", "int", "", "是", "测评分数"),
            ("risk_probability", "decimal", "5,2", "是", "风险概率"),
            ("result_level", "varchar", "32", "是", "结果等级"),
            ("analysis", "text", "", "是", "分析说明"),
            ("model_name", "varchar", "64", "是", "模型名称"),
            ("leading_factors_json", "json", "", "是", "主要影响因素"),
            ("status", "varchar", "32", "是", "记录状态"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-7 聊天消息表结构",
        "name": "chat_message",
        "cn": "聊天消息表",
        "desc": "聊天消息表保存用户与心理咨询师之间的在线沟通记录，包括发送者、接收者、消息内容、附件地址、敏感标记和审核状态。该表使即时通信不仅能在页面端实时展示，也能在后端形成可追溯的数据记录。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("sender_id", "bigint", "", "否", "发送者编号"),
            ("receiver_id", "bigint", "", "否", "接收者编号"),
            ("content", "text", "", "是", "消息内容"),
            ("file_url", "varchar", "255", "是", "附件地址"),
            ("sensitive_flag", "tinyint", "", "是", "敏感标记"),
            ("review_status", "varchar", "32", "是", "审核状态"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-8 操作日志表结构",
        "name": "operation_log",
        "cn": "操作日志表",
        "desc": "操作日志表用于记录用户或管理员在系统中的关键操作，包括操作模块、动作类型、来源 IP 和详细内容。该表可以用于问题排查、后台审计和异常行为追踪，能够增强系统运行过程的可管理性。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("user_id", "bigint", "", "是", "操作用户编号"),
            ("module", "varchar", "64", "是", "操作模块"),
            ("action", "varchar", "64", "是", "操作动作"),
            ("ip", "varchar", "64", "是", "来源 IP"),
            ("detail", "text", "", "是", "操作详情"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-9 心理文章表结构",
        "name": "article",
        "cn": "心理文章表",
        "desc": "心理文章表保存平台心理科普内容，包括文章标题、分类、摘要、正文、作者和发布状态。该表支撑用户端心理文章浏览和管理员端文章维护，正文使用 LONGTEXT 类型，能够保存较长篇幅的心理健康科普材料。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("title", "varchar", "255", "否", "文章标题"),
            ("category", "varchar", "64", "是", "文章分类"),
            ("summary", "text", "", "是", "文章摘要"),
            ("content", "longtext", "", "否", "文章正文"),
            ("author_name", "varchar", "64", "是", "作者名称"),
            ("status", "varchar", "32", "是", "发布状态"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
    {
        "caption": "表4-10 通知公告表结构",
        "name": "notification",
        "cn": "通知公告表",
        "desc": "通知公告表保存系统通知和管理员公告，支持按指定用户或指定角色进行消息投放。该表能够覆盖账号状态变更提醒、平台公告、预约相关通知等场景，使用户端和后台端的信息传递更加统一。",
        "rows": [
            ("id", "bigint", "", "否", "主键"),
            ("target_user_id", "bigint", "", "是", "目标用户编号"),
            ("target_role", "varchar", "32", "是", "目标角色"),
            ("title", "varchar", "255", "是", "通知标题"),
            ("content", "text", "", "是", "通知内容"),
            ("created_by", "bigint", "", "是", "创建人编号"),
            ("created_at", "datetime", "", "是", "创建时间"),
            ("updated_at", "datetime", "", "是", "更新时间"),
            ("deleted", "tinyint", "", "是", "逻辑删除标记"),
        ],
    },
]


def get_font(size: int) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text, font, fill="black"):
    x1, y1, x2, y2 = box
    lines = []
    for part in str(text).split("\n"):
        if not part:
            lines.append("")
            continue
        line = ""
        for ch in part:
            test = line + ch
            if draw.textbbox((0, 0), test, font=font)[2] <= (x2 - x1 - 18):
                line = test
            else:
                lines.append(line)
                line = ch
        lines.append(line)
    line_heights = [draw.textbbox((0, 0), line or "口", font=font)[3] - draw.textbbox((0, 0), line or "口", font=font)[1] for line in lines]
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 5
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += h + 5


def rect(draw, center, size, label, font):
    x, y = center
    w, h = size
    box = (x - w / 2, y - h / 2, x + w / 2, y + h / 2)
    draw.rectangle(box, outline="black", width=3)
    draw_centered_text(draw, box, label, font)
    return box


def oval(draw, center, size, label, font):
    x, y = center
    w, h = size
    box = (x - w / 2, y - h / 2, x + w / 2, y + h / 2)
    draw.ellipse(box, outline="black", width=2)
    draw_centered_text(draw, box, label, font)
    return box


def diamond(draw, center, size, label, font):
    x, y = center
    w, h = size
    points = [(x, y - h / 2), (x + w / 2, y), (x, y + h / 2), (x - w / 2, y)]
    draw.polygon(points, outline="black", fill="white")
    draw.line(points + [points[0]], fill="black", width=3)
    draw_centered_text(draw, (x - w / 2, y - h / 2, x + w / 2, y + h / 2), label, font)
    return (x - w / 2, y - h / 2, x + w / 2, y + h / 2)


def line(draw, a, b, label=None, font=None, offset=(0, 0)):
    draw.line([a, b], fill="black", width=2)
    if label and font:
        mx, my = (a[0] + b[0]) / 2 + offset[0], (a[1] + b[1]) / 2 + offset[1]
        draw.text((mx, my), label, font=font, fill="black")


def polyline(draw, points, label=None, font=None, label_at=0.5, offset=(0, 0)):
    draw.line(points, fill="black", width=2)
    if label and font:
        idx = max(0, min(len(points) - 2, int((len(points) - 1) * label_at)))
        a, b = points[idx], points[idx + 1]
        mx, my = (a[0] + b[0]) / 2 + offset[0], (a[1] + b[1]) / 2 + offset[1]
        draw.text((mx, my), label, font=font, fill="black")


def generate_er_image():
    img = Image.new("RGB", (2400, 1550), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(46)
    entity_font = get_font(34)
    small_font = get_font(27)
    attr_font = get_font(25)

    draw.text((860, 35), "心理健康服务平台数据库 E-R 图", font=title_font, fill="black")

    entities = {
        "user": (260, 720, "用户"),
        "profile": (760, 270, "咨询师资料"),
        "appointment": (760, 560, "预约记录"),
        "review": (760, 850, "咨询评价"),
        "scale": (1400, 270, "测评量表"),
        "record": (1400, 560, "测评记录"),
        "chat": (1400, 850, "聊天消息"),
        "article": (1980, 270, "心理文章"),
        "notification": (1980, 560, "通知公告"),
        "log": (1980, 850, "操作日志"),
    }
    for _, (x, y, label) in entities.items():
        rect(draw, (x, y), (210, 80), label, entity_font)

    # central user attributes
    oval(draw, (105, 540), (150, 62), "用户编号", attr_font)
    oval(draw, (330, 535), (150, 62), "用户名", attr_font)
    oval(draw, (105, 920), (150, 62), "角色", attr_font)
    oval(draw, (330, 925), (150, 62), "账号状态", attr_font)
    line(draw, (160, 565), (220, 685))
    line(draw, (330, 566), (285, 685))
    line(draw, (160, 895), (220, 755))
    line(draw, (325, 895), (285, 755))

    # relationship diamonds in lanes; layout avoids crossing lines.
    diamond(draw, (500, 270), (130, 90), "拥有", small_font)
    polyline(draw, [(365, 690), (430, 690), (430, 270), (435, 270)], "1", small_font, 0.8, (-28, -35))
    line(draw, (565, 270), (655, 270), "1", small_font, (10, -35))

    diamond(draw, (500, 560), (130, 90), "预约", small_font)
    polyline(draw, [(365, 710), (430, 710), (430, 560), (435, 560)], "1", small_font, 0.8, (-28, -35))
    line(draw, (565, 560), (655, 560), "n", small_font, (10, -35))

    diamond(draw, (500, 850), (130, 90), "评价", small_font)
    polyline(draw, [(365, 735), (430, 735), (430, 850), (435, 850)], "1", small_font, 0.8, (-28, 10))
    line(draw, (565, 850), (655, 850), "n", small_font, (10, -35))

    diamond(draw, (1080, 270), (150, 90), "配置", small_font)
    line(draw, (865, 270), (1005, 270), "1", small_font, (-15, -35))
    line(draw, (1155, 270), (1295, 270), "n", small_font, (10, -35))

    diamond(draw, (1080, 560), (150, 90), "生成", small_font)
    polyline(draw, [(365, 720), (965, 720), (965, 560), (1005, 560)], "1", small_font, 0.2, (180, -35))
    line(draw, (1155, 560), (1295, 560), "n", small_font, (10, -35))
    diamond(draw, (1400, 415), (150, 90), "使用", small_font)
    line(draw, (1400, 310), (1400, 370), "1", small_font, (12, -20))
    line(draw, (1400, 460), (1400, 520), "n", small_font, (12, 5))

    diamond(draw, (1080, 850), (150, 90), "收发", small_font)
    polyline(draw, [(365, 740), (965, 740), (965, 850), (1005, 850)], "1", small_font, 0.2, (180, 10))
    line(draw, (1155, 850), (1295, 850), "n", small_font, (10, -35))

    diamond(draw, (1690, 270), (150, 90), "维护", small_font)
    polyline(draw, [(365, 700), (1560, 700), (1560, 270), (1615, 270)], "1", small_font, 0.2, (430, -35))
    line(draw, (1765, 270), (1875, 270), "n", small_font, (8, -35))

    diamond(draw, (1690, 560), (150, 90), "发布", small_font)
    polyline(draw, [(365, 720), (1580, 720), (1580, 560), (1615, 560)], "1", small_font, 0.2, (450, -35))
    line(draw, (1765, 560), (1875, 560), "n", small_font, (8, -35))

    diamond(draw, (1690, 850), (150, 90), "记录", small_font)
    polyline(draw, [(365, 740), (1580, 740), (1580, 850), (1615, 850)], "1", small_font, 0.2, (450, 8))
    line(draw, (1765, 850), (1875, 850), "n", small_font, (8, -35))

    # selected attributes for major entities
    attrs = [
        ((760, 140), "擅长方向", "profile"), ((940, 220), "职称", "profile"), ((940, 335), "在线状态", "profile"), ((760, 405), "评分", "profile"),
        ((760, 455), "预约时间", "appointment"), ((945, 560), "预约状态", "appointment"), ((760, 665), "咨询方式", "appointment"),
        ((760, 745), "评分", "review"), ((945, 850), "评价内容", "review"),
        ((1400, 140), "量表名称", "scale"), ((1585, 270), "量表编码", "scale"), ((1400, 375), "题目配置", "scale"),
        ((1400, 455), "答案数据", "record"), ((1585, 560), "风险概率", "record"), ((1400, 665), "结果等级", "record"),
        ((1400, 745), "消息内容", "chat"), ((1585, 850), "审核状态", "chat"),
        ((1980, 155), "文章标题", "article"), ((2180, 270), "发布状态", "article"),
        ((1980, 445), "通知标题", "notification"), ((2180, 560), "目标角色", "notification"),
        ((1980, 735), "操作模块", "log"), ((2180, 850), "操作详情", "log"),
    ]
    entity_points = {k: (x, y) for k, (x, y, _) in entities.items()}
    for (x, y), label, key in attrs:
        oval(draw, (x, y), (165, 58), label, attr_font)
        ex, ey = entity_points[key]
        line(draw, (x, y + (29 if y < ey else -29)), (ex, ey - 40 if y < ey else ey + 40))

    # Profile also participates in appointment/review through straight local relations.
    diamond(draw, (760, 415), (130, 80), "被预约", small_font)
    line(draw, (760, 310), (760, 375), "1", small_font, (12, -20))
    line(draw, (760, 455), (760, 520), "n", small_font, (12, 5))
    diamond(draw, (930, 710), (130, 80), "被评价", small_font)
    polyline(draw, [(825, 295), (930, 295), (930, 670)], "1", small_font, 0.6, (10, 120))
    polyline(draw, [(930, 750), (930, 850), (865, 850)], "n", small_font, 0.2, (14, -10))

    img.save(ER_IMAGE)


def add_run_text(paragraph, text, font_size=10.5, bold=False):
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)
    run.bold = bold
    return run


def set_paragraph_format(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)


def insert_paragraph_after(element, text="", style=None, align=None, first_line=True):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, DOC_PARENT)
    if style:
        paragraph.style = style
    if align is not None:
        paragraph.alignment = align
    if text:
        add_run_text(paragraph, text)
    set_paragraph_format(paragraph, first_line=first_line)
    return paragraph


def move_table_after(doc, element, table):
    element.addnext(table._tbl)
    return table._tbl


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(font_size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8" if edge in ("top", "bottom") else "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def clear_table_style(table):
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblStyle", "w:tblLook"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)


def set_table_layout(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_schema_table(doc, after_element, table_info):
    caption = insert_paragraph_after(after_element, table_info["caption"], align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    caption.runs[0].font.size = Pt(11)
    caption.runs[0].bold = False
    after_element = caption._p

    rows = table_info["rows"]
    table = doc.add_table(rows=len(rows) + 1, cols=6)
    move_table_after(doc, after_element, table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_style(table)
    set_table_borders(table)
    headers = ["序号", "列名", "数据类型", "长度", "允许空", "说明"]
    widths = [760, 1900, 1300, 900, 900, 3000]
    set_table_layout(table, widths)
    for col_idx, header in enumerate(headers):
        set_cell_text(table.cell(0, col_idx), header, bold=True, font_size=9.5)
        set_cell_margins(table.cell(0, col_idx))
    for row_idx, row in enumerate(rows, 1):
        values = [row_idx] + list(row)
        for col_idx, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (1, 5) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(row_idx, col_idx), value, align=align, font_size=8.5)
            set_cell_margins(table.cell(row_idx, col_idx), top=60, bottom=60)
    after_element = table._tbl

    paragraph = insert_paragraph_after(after_element, table_info["desc"], first_line=True)
    after_element = paragraph._p
    return after_element


def rebuild_database_section():
    global DOC_PARENT
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(INPUT_DOCX)
    shutil.copy2(INPUT_DOCX, WORK_DOCX)
    doc = Document(str(WORK_DOCX))
    DOC_PARENT = doc._body

    db_heading = None
    next_heading = None
    seen = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "数据库设计":
            db_heading = p
            seen = True
            continue
        if seen and p.style.name.startswith("Heading 2"):
            next_heading = p
            break
    if db_heading is None or next_heading is None:
        raise RuntimeError("Could not locate database section boundaries")

    body = doc._body._element
    children = list(body)
    start = children.index(db_heading._p)
    end = children.index(next_heading._p)
    for child in children[start + 1:end]:
        body.remove(child)

    after = db_heading._p
    intro = (
        "数据库设计围绕用户、咨询师、预约、测评、在线沟通和后台管理等业务对象展开。"
        "其中，用户表保存平台统一身份信息，咨询师资料表用于扩展咨询师业务属性，预约记录表连接普通用户与心理咨询师，"
        "测评量表表和测评记录表支撑心理测评功能，聊天消息表保存在线咨询过程，文章表和通知公告表负责平台内容与消息发布，"
        "操作日志表则记录系统关键操作。通过这样的设计，系统能够把账号身份、业务过程和统计分析数据区分存储，同时又能通过编号关联形成完整的数据链路。"
    )
    p = insert_paragraph_after(after, intro)
    after = p._p

    p = insert_paragraph_after(after, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    run = p.add_run()
    run.add_picture(str(ER_IMAGE), width=Cm(15.2))
    after = p._p
    cap = insert_paragraph_after(after, "图4-6 数据库 E-R 图", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    after = cap._p
    explain = (
        "如图4-6所示，系统数据库以用户实体为中心向外展开。一个用户账号可以对应一条咨询师资料，也可以发起多条预约记录、提交多条测评记录、发送或接收多条聊天消息；"
        "咨询师资料与预约记录、咨询评价之间形成一对多关系，测评量表与测评记录之间也形成一对多关系；管理员或系统账号可以维护心理文章、发布通知公告，并在操作日志中留下处理记录。"
        "图中矩形表示实体，菱形表示实体之间的联系，椭圆表示主要属性，连线旁的一和 n 用来说明一对一或一对多的业务关系。"
    )
    p = insert_paragraph_after(after, explain)
    after = p._p

    p = insert_paragraph_after(after, "根据数据库建表语句和后端初始化逻辑，平台主要包含以下数据表。各表字段结构如下所示。")
    after = p._p
    for info in TABLES:
        after = add_schema_table(doc, after, info)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    generate_er_image()
    rebuild_database_section()
    print(f"ER image written: {ER_IMAGE}")
    print(f"DOCX written: {OUTPUT_DOCX}")
