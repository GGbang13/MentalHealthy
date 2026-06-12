from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\桌面\毕业相关\附件6-2：毕业论文&开题报告模板.docx")
FALLBACK_TEMPLATE = ROOT / "project_previous.docx"
SOURCE_MD = ROOT / "frontend" / "docs" / "thesis-draft.md"
ASSET_DIR = ROOT / "docs" / "thesis-assets"
OUTPUT = ROOT / "毕业论文初稿-按附件6-2模板-扩展版.docx"


@dataclass
class Block:
    kind: str
    level: int
    text: str


def parse_markdown(path: Path) -> list[Block]:
    blocks: list[Block] = []
    pending: list[str] = []

    def flush() -> None:
        nonlocal pending
        if pending:
            blocks.append(Block("paragraph", 0, " ".join(pending).strip()))
            pending = []

    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if match:
            flush()
            blocks.append(Block("heading", len(match.group(1)), match.group(2).strip()))
            continue
        if re.match(r"^(\[\d+\]|\d+[.)、．])\s*", stripped):
            flush()
            blocks.append(Block("paragraph", 0, stripped))
            continue
        pending.append(stripped)
    flush()
    return blocks


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\s*", "", text).strip()


def split_sections(blocks: list[Block]) -> dict[str, list[Block]]:
    sections = {"abstract_zh": [], "abstract_en": [], "main": [], "references": []}
    target = "ignore"
    for block in blocks:
        if block.kind == "heading" and block.level == 1:
            continue
        if block.kind == "heading" and block.level == 2:
            title = block.text.strip()
            if title == "摘要":
                target = "abstract_zh"
                continue
            if title == "Abstract":
                target = "abstract_en"
                continue
            if title in {"目录", "浙大宁波理工学院本科毕业论文（设计）承诺书", "致谢"}:
                target = "ignore"
                continue
            if title in {"参考文献", "7 参考文献"}:
                target = "references"
                sections[target].append(block)
                continue
            if re.match(r"^\d+\s+", title):
                target = "main"
        if target != "ignore":
            sections[target].append(block)
    return sections


def title_from(blocks: list[Block]) -> str:
    return next(block.text for block in blocks if block.kind == "heading" and block.level == 1)


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_paragraph(paragraph, first_line: bool = False, align=None, spacing: float = 1.5) -> None:
    paragraph.paragraph_format.line_spacing = spacing
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, 12)


def add_para(document: Document, text: str = "", *, first_line: bool = True, align=None, bold: bool = False, size: int = 12):
    paragraph = document.add_paragraph()
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size, bold)
    format_paragraph(paragraph, first_line=first_line, align=align)
    return paragraph


def add_heading(document: Document, title: str, level: int) -> None:
    if level == 1:
        display, size = title, 16
    elif level == 2:
        display, size = title, 15
    else:
        display, size = title, 14
    paragraph = document.add_paragraph()
    run = paragraph.add_run(display)
    set_run_font(run, size, True)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    format_paragraph(paragraph, first_line=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, 10, False)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(6)


def add_picture(document: Document, image_path: Path, caption: str, width: float = 5.8) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(document, caption)


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, 10, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, 9)
    document.add_paragraph()


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = ""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, 10)


def add_cover(document: Document, title: str) -> None:
    add_para(document, "公开论文", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    document.add_paragraph()
    add_para(document, "毕业论文（设计）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    document.add_paragraph()
    add_para(document, title, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    document.add_paragraph()
    for line in [
        f"题    目    {title}",
        "姓    名    [待填写]",
        "学    号    [待填写]",
        "专业班级    [待填写]",
        "指导教师    [待填写]",
        "学    院    [待填写]",
        "日    期    2026 年 5 月 1 日",
    ]:
        add_para(document, line, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    document.add_page_break()


def add_commitment(document: Document) -> None:
    add_para(document, "浙大宁波理工学院本科毕业论文（设计）承诺书", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    items = [
        "1.本人郑重承诺所呈交的毕业论文（设计），是在指导教师的指导下严格按照学校和学院有关规定完成的。",
        "2.本人在毕业论文（设计）中引用他人的观点和参考资料均加以注释和说明。",
        "3.与我一同工作过的同学对本研究所做的任何贡献均已在论文中做了明确的说明并表示谢意。",
        "4.本人承诺在毕业论文（设计）工作过程中没有抄袭他人研究成果和伪造数据等行为。",
        "5.若本人在毕业论文（设计）中有任何侵犯知识产权的行为，由本人承担相应的法律责任。",
        "6.本人完全了解浙大宁波理工学院有权保留并向有关部门或机构送交本论文（设计）复印件和电子文档，允许本论文（设计）被查阅和借阅。",
    ]
    for item in items:
        add_para(document, item)
    document.add_paragraph()
    add_para(document, "作者签名：                              日期：2026 年 5 月 1 日", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    document.add_page_break()


def add_acknowledgement(document: Document) -> None:
    add_para(document, "致谢", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    for text in [
        "在本课题完成过程中，首先感谢指导教师在选题论证、系统设计、功能实现和论文写作等阶段给予的指导与建议。老师对系统业务边界、技术路线和论文结构提出了许多具体意见，使本文能够围绕实际项目展开分析与总结。",
        "同时感谢学院各位老师在本科阶段的课程教学和实践训练中给予的帮助。软件工程、数据库原理、Web 前端开发、Java 程序设计等课程为本系统的设计与实现奠定了基础。",
        "感谢同学和朋友在系统测试、页面体验反馈和论文修改过程中提供的帮助。最后感谢家人长期以来的支持与鼓励，使我能够以稳定的状态完成毕业设计相关工作。",
    ]:
        add_para(document, text)
    document.add_page_break()


def create_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    font_path = "C:/Windows/Fonts/msyh.ttc"
    title_font = ImageFont.truetype(font_path, 30)
    font = ImageFont.truetype(font_path, 22)
    small = ImageFont.truetype(font_path, 18)

    def box(draw, xy, text, fill, outline="#24584A"):
        draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
        lines = text.split("\n")
        y = xy[1] + 20
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((xy[0] + (xy[2] - xy[0] - bbox[2]) / 2, y), line, font=font, fill="#18352F")
            y += 32

    def arrow(draw, start, end, color="#24584A"):
        draw.line((*start, *end), fill=color, width=4)
        dx = end[0] - start[0]
        dy = end[1] - start[1]
        if abs(dx) >= abs(dy):
            if dx >= 0:
                points = [(end[0], end[1]), (end[0] - 18, end[1] - 10), (end[0] - 18, end[1] + 10)]
            else:
                points = [(end[0], end[1]), (end[0] + 18, end[1] - 10), (end[0] + 18, end[1] + 10)]
        else:
            if dy >= 0:
                points = [(end[0], end[1]), (end[0] - 10, end[1] - 18), (end[0] + 10, end[1] - 18)]
            else:
                points = [(end[0], end[1]), (end[0] - 10, end[1] + 18), (end[0] + 10, end[1] + 18)]
        draw.polygon(points, fill=color)

    img = Image.new("RGB", (1500, 850), "#F7F4EC")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "心理健康服务平台系统架构图", font=title_font, fill="#18352F")
    layers = [
        ("用户访问层\n浏览器 / 用户端 / 咨询师端 / 管理员端", "#E8F5F2"),
        ("前端应用层\nVue 3 / Router / Pinia / Element Plus / ECharts", "#EAF1FF"),
        ("接口通信层\nAxios REST API / JWT / WebSocket", "#FFF3D6"),
        ("后端业务层\nSpring Boot / Security / Controller / Service", "#E9F7E9"),
        ("数据与模型层\nMySQL / Redis / CatBoost 模型 / 操作日志", "#FCE9E2"),
    ]
    y = 110
    for text, fill in layers:
        box(draw, (160, y, 1340, y + 100), text, fill)
        if y < 590:
            arrow(draw, (750, y + 100), (750, y + 140))
        y += 135
    img.save(ASSET_DIR / "architecture.png")

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "心理健康服务平台业务流程图", font=title_font, fill="#18352F")
    flow_nodes = [
        ((70, 150, 330, 250), "用户注册/登录", "#E8F5F2"),
        ((460, 150, 720, 250), "阅读文章\n完成心理测评", "#EAF1FF"),
        ((850, 150, 1110, 250), "选择咨询师\n提交预约", "#FFF3D6"),
        ((1240, 150, 1500, 250), "咨询师审核\n确认/拒绝", "#E9F7E9"),
        ((1240, 440, 1500, 540), "已确认预约\n进入聊天", "#FCE9E2"),
        ((850, 440, 1110, 540), "消息持久化\nWebSocket 推送", "#EAF1FF"),
        ((460, 440, 720, 540), "测评/预约/聊天\n形成服务记录", "#F4FAF8"),
        ((70, 440, 330, 540), "管理员看板\n风险监控与内容治理", "#FFF3D6"),
    ]
    for xy, text, fill in flow_nodes:
        box(draw, xy, text, fill)
    for start, end in [
        ((330, 200), (460, 200)),
        ((720, 200), (850, 200)),
        ((1110, 200), (1240, 200)),
        ((1370, 250), (1370, 440)),
        ((1240, 490), (1110, 490)),
        ((850, 490), (720, 490)),
        ((460, 490), (330, 490)),
    ]:
        arrow(draw, start, end)
    draw.text((870, 315), "拒绝时返回预约列表并保留状态记录", font=small, fill="#7B4A12")
    arrow(draw, (1240, 235), (1110, 440), "#9A6A1F")
    img.save(ASSET_DIR / "business-flow.png")

    img = Image.new("RGB", (1600, 1000), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "心理健康服务平台 E-R 图", font=title_font, fill="#18352F")
    entities = {
        "sys_user\n用户/角色/资料": (80, 130, 390, 250),
        "counselor_profile\n咨询师资料": (620, 130, 930, 250),
        "appointment\n预约记录": (1140, 130, 1450, 250),
        "assessment_scale\n测评量表": (80, 470, 390, 590),
        "assessment_record\n测评记录": (620, 470, 930, 590),
        "chat_message\n聊天消息": (1140, 470, 1450, 590),
        "article\n心理文章": (350, 770, 660, 890),
        "notification\n通知信息": (900, 770, 1210, 890),
    }
    for text, xy in entities.items():
        box(draw, xy, text, "#F4FAF8")
    for start, end, label in [
        ((390, 190), (620, 190), "1:1"),
        ((930, 190), (1140, 190), "1:N"),
        ((235, 250), (775, 470), "1:N"),
        ((390, 530), (620, 530), "1:N"),
        ((235, 250), (1140, 530), "1:N"),
        ((235, 250), (505, 770), "1:N"),
        ((235, 250), (1055, 770), "1:N"),
    ]:
        draw.line((*start, *end), fill="#24584A", width=4)
        mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
        draw.text((mx - 24, my - 24), label, font=small, fill="#24584A")
    img.save(ASSET_DIR / "er-diagram.png")

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "心理测评风险分析流程图", font=title_font, fill="#18352F")
    assessment_nodes = [
        ((110, 140, 420, 240), "读取量表配置\nquestion_json", "#E8F5F2"),
        ((560, 140, 870, 240), "前端动态渲染\n题目与选项", "#EAF1FF"),
        ((1010, 140, 1320, 240), "用户提交答案\nanswer_json", "#FFF3D6"),
        ((1010, 390, 1320, 490), "特征解析\n分数计算", "#F4FAF8"),
        ((560, 390, 870, 490), "CatBoost 推理\n或 Java 回退", "#FCE9E2"),
        ((110, 390, 420, 490), "生成风险等级\n概率与影响因素", "#E9F7E9"),
        ((560, 640, 870, 740), "结果落库\n历史记录/监控", "#FFF3D6"),
    ]
    for xy, text, fill in assessment_nodes:
        box(draw, xy, text, fill)
    for start, end in [
        ((420, 190), (560, 190)),
        ((870, 190), (1010, 190)),
        ((1165, 240), (1165, 390)),
        ((1010, 440), (870, 440)),
        ((560, 440), (420, 440)),
        ((265, 490), (560, 690)),
    ]:
        arrow(draw, start, end)
    img.save(ASSET_DIR / "assessment-flow.png")

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 40), "在线聊天模块时序图", font=title_font, fill="#18352F")
    actors = [("用户端", 160), ("后端服务", 570), ("数据库", 980), ("咨询师端", 1390)]
    for name, x in actors:
        draw.text((x - 45, 115), name, font=font, fill="#18352F")
        draw.line((x, 160, x, 790), fill="#8AA99D", width=3)
    steps = [
        (210, 160, 570, "携带 JWT 建立连接"),
        (280, 570, 980, "校验身份并注册会话"),
        (360, 160, 570, "REST 发送消息"),
        (430, 570, 980, "写入 chat_message"),
        (500, 980, 570, "返回持久化消息"),
        (570, 570, 1390, "WebSocket 推送"),
        (650, 570, 160, "回推发送结果"),
        (720, 1390, 570, "回复消息并复用同一链路"),
    ]
    for y, x1, x2, label in steps:
        arrow(draw, (x1, y), (x2, y))
        draw.text((min(x1, x2) + 35, y - 28), label, font=small, fill="#18352F")
    img.save(ASSET_DIR / "chat-sequence.png")

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "表3.1  核心数据表设计", font=title_font, fill="#18352F")
    headers = ["数据表", "主要字段", "功能说明"]
    rows = [
        ["sys_user", "username、password、role、email、phone、status", "保存用户账号、角色和基础资料，是登录认证与权限判断的数据基础。"],
        ["counselor_profile", "user_id、title、specialties、price_per_hour、rating", "保存咨询师执业资料，与用户表形成一对一关联。"],
        ["appointment", "user_id、counselor_id、appointment_time、status", "保存预约申请和状态流转，是咨询服务闭环的核心表。"],
        ["assessment_scale", "name、code、question_json、rule_json、enabled", "保存量表定义和题目配置，支持动态渲染测评页面。"],
        ["assessment_record", "user_id、scale_id、score、risk_probability、result_level", "保存用户测评结果、风险等级和模型解释信息。"],
        ["chat_message", "sender_id、receiver_id、content、file_url、review_status", "保存用户与咨询师之间的聊天消息和审核状态。"],
        ["article、notification", "title、content、status、target_role、created_at", "支撑心理文章发布、平台通知和内容运营。"],
    ]
    x_positions = [60, 330, 850, 1540]
    y = 110
    row_h = 88
    draw.rectangle((60, y, 1540, y + row_h), fill="#E8F5F2", outline="#24584A", width=3)
    for i, header in enumerate(headers):
        draw.text((x_positions[i] + 18, y + 28), header, font=font, fill="#18352F")
    y += row_h
    for index, row in enumerate(rows):
        fill = "#FFFFFF" if index % 2 == 0 else "#F7FAF9"
        draw.rectangle((60, y, 1540, y + row_h), fill=fill, outline="#24584A", width=2)
        for x in x_positions[1:-1]:
            draw.line((x, y, x, y + row_h), fill="#24584A", width=2)
        for i, value in enumerate(row):
            max_chars = 16 if i == 0 else 27 if i == 1 else 34
            lines = [value[j:j + max_chars] for j in range(0, len(value), max_chars)]
            for line_no, line in enumerate(lines[:3]):
                draw.text((x_positions[i] + 18, y + 14 + line_no * 24), line, font=small, fill="#18352F")
        y += row_h
    img.save(ASSET_DIR / "database-table.png")

    img = Image.new("RGB", (1600, 900), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.text((60, 35), "表6.1  后续增加工作量方向", font=title_font, fill="#18352F")
    headers = ["方向", "具体工作", "预期价值"]
    rows = [
        ["移动端与小程序", "补充移动端专项页面、微信小程序入口、消息订阅与端侧适配测试。", "提升学生高频使用场景的触达率。"],
        ["危机干预闭环", "增加高风险工单、人工复核、转介记录、院系协同与处置时限跟踪。", "增强平台在真实心理服务制度中的可用性。"],
        ["模型治理", "补充训练数据版本、模型评估指标、解释一致性审查、灰度发布与回滚机制。", "降低算法误用风险，提升结果可信度。"],
        ["审计与合规", "增加管理员操作日志、敏感数据导出审批、数据留存策略与脱敏统计接口。", "满足长期部署中的隐私保护和责任追溯要求。"],
        ["音视频咨询", "接入实时音视频、会前知情同意、会后咨询记录模板与会话质量评价。", "扩展服务形态，提升在线咨询连续性。"],
        ["运营分析", "增加预约趋势、咨询师工作量、文章阅读转化、复测率和服务满意度分析。", "为心理中心资源配置和内容运营提供依据。"],
    ]
    x_positions = [60, 330, 930, 1540]
    y = 110
    row_h = 98
    draw.rectangle((60, y, 1540, y + row_h), fill="#E8F5F2", outline="#24584A", width=3)
    for i, header in enumerate(headers):
        draw.text((x_positions[i] + 18, y + 34), header, font=font, fill="#18352F")
    y += row_h
    for index, row in enumerate(rows):
        fill = "#FFFFFF" if index % 2 == 0 else "#F7FAF9"
        draw.rectangle((60, y, 1540, y + row_h), fill=fill, outline="#24584A", width=2)
        for x in x_positions[1:-1]:
            draw.line((x, y, x, y + row_h), fill="#24584A", width=2)
        for i, value in enumerate(row):
            max_chars = 9 if i == 0 else 28
            lines = [value[j:j + max_chars] for j in range(0, len(value), max_chars)]
            for line_no, line in enumerate(lines[:3]):
                draw.text((x_positions[i] + 18, y + 16 + line_no * 26), line, font=small, fill="#18352F")
        y += row_h
    img.save(ASSET_DIR / "workload-table.png")


def add_abstract(document: Document, zh: list[Block], en: list[Block]) -> None:
    add_para(document, "摘要", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    for block in zh:
        if block.kind == "paragraph":
            add_para(document, block.text)
    document.add_page_break()
    add_para(document, "Abstract", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    for block in en:
        if block.kind == "paragraph":
            add_para(document, block.text)
    document.add_page_break()


def add_toc_placeholder(document: Document) -> None:
    add_para(document, "目录", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
    toc_items = [
        ("1  绪论", "7", 0),
        ("1.1  研究背景与意义", "7", 1),
        ("1.2  国内外研究现状", "7", 1),
        ("1.3  论文的主要内容和组织结构", "8", 1),
        ("1.4  技术路线", "8", 1),
        ("2  系统需求分析", "9", 0),
        ("2.1  可行性分析", "9", 1),
        ("2.2  角色与业务需求", "9", 1),
        ("2.3  功能需求", "9", 1),
        ("2.4  非功能需求", "10", 1),
        ("3  系统总体设计", "10", 0),
        ("3.1  系统架构设计", "10", 1),
        ("3.2  前端工程设计", "11", 1),
        ("3.3  后端接口与数据交互设计", "11", 1),
        ("3.4  数据库概念结构设计", "12", 1),
        ("4  系统核心功能实现", "13", 0),
        ("4.1  登录认证与权限控制", "13", 1),
        ("4.2  用户端功能实现", "14", 1),
        ("4.3  咨询师端功能实现", "15", 1),
        ("4.4  管理员端功能实现", "16", 1),
        ("4.5  心理测评与风险分析功能实现", "17", 1),
        ("4.6  在线聊天功能实现", "18", 1),
        ("5  系统测试", "20", 0),
        ("5.1  测试环境", "20", 1),
        ("5.2  功能测试", "20", 1),
        ("5.3  兼容性与异常测试", "21", 1),
        ("5.4  测试结果分析", "21", 1),
        ("6  总结与展望", "21", 0),
        ("6.1  项目总结", "21", 1),
        ("6.2  不足与展望", "22", 1),
        ("7  参考文献", "22", 0),
    ]

    def display_width(value: str) -> int:
        width = 0
        for char in value:
            width += 1 if ord(char) < 128 else 2
        return width

    for title, page, level in toc_items:
        title_text = ("    " if level else "") + title
        dot_count = max(6, 84 - display_width(title_text) - len(page))
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(0)
        run_title = paragraph.add_run(title_text + " ")
        set_run_font(run_title, 12, False)
        run_dots = paragraph.add_run("." * dot_count)
        set_run_font(run_dots, 12, False)
        run_page = paragraph.add_run(" " + page)
        set_run_font(run_page, 12, False)
    document.add_page_break()


def expansion_after_heading(title: str) -> list[str]:
    extras = {
        "登录认证与权限控制": [
            "后端登录流程由认证控制器接收用户名和密码，业务层根据用户名查询用户记录，并使用 BCrypt 对密码摘要进行校验。校验通过后，系统根据用户编号和角色生成 JWT，前端获得 token 后写入 Pinia 状态和 localStorage。该流程将会话状态放在客户端保存，服务端接口通过过滤器解析请求头中的身份令牌，从而避免传统 Session 在前后端分离场景中的状态同步问题。",
            "权限控制采用前端路由守卫和后端接口校验共同完成。前端在路由 meta 中配置 allowedRoles，用户访问页面前先判断是否登录以及当前角色是否具备访问权限；后端则在业务接口中根据认证信息判断当前用户是否允许操作目标数据。例如普通用户只能取消自己的预约，咨询师只能处理属于自己的预约，聊天发送前需要检查双方是否存在已确认预约或历史会话。",
        ],
        "用户端功能实现": [
            "用户端页面强调服务入口的集中展示。首页将心理测评、预约咨询、心理文章和在线沟通作为四个主要卡片入口，并配合通知中心展示当前角色可见的通知信息。该设计让普通用户进入平台后能够快速定位核心功能，减少在管理菜单中反复查找的成本。",
            "预约流程由咨询师列表和预约表单共同组成。用户浏览咨询师时可查看职称、擅长方向、从业年限、咨询价格、在线状态和评分等信息。进入预约页面后，系统根据表单字段收集咨询师、预约时间、咨询形式、咨询时长和问题描述。预约提交后记录进入待确认状态，等待咨询师处理。",
        ],
        "心理测评与风险分析功能实现": [
            "测评模块的量表问题使用 JSON 保存，使题目内容、取值范围、步长和是否反向等配置能够由后端统一维护。前端解析 questionJson 后动态生成表单控件，因此后续新增 PHQ-9、GAD-7 或其他心理量表时，不需要大幅修改页面结构。",
            "风险分析结果由分数、风险概率、风险等级、模型名称和主要影响因素组成。若 CatBoost 模型文件存在，后端可优先调用 Python 推理脚本生成预测结果；若模型不可用，则回退到规则或权重方式进行基础评估。这种设计兼顾了演示环境稳定性和后续接入真实训练模型的扩展空间。",
            "主要影响因素以 leadingFactors 形式返回给前端，每个因素包含名称、得分、影响度、方向和说明。页面可将风险项与保护项分开展示，使用户和管理员不只看到一个分数，还能理解哪些变量对结果贡献较大。需要说明的是，该贡献反映模型预测的重要性，不等同于医学因果判断。",
        ],
        "在线聊天功能实现": [
            "聊天模块采用 REST API 与 WebSocket 结合的方式实现。页面加载时先通过联系人接口获取可沟通对象，选择联系人后调用历史消息接口按时间顺序加载消息；发送消息时通过接口保存消息并返回持久化后的记录。这样即使实时连接短暂不可用，系统仍能完成基本消息收发。",
            "WebSocket 连接用于提升沟通及时性。前端在连接地址中携带 token，握手拦截器解析身份后登记当前用户会话。收到新消息时，若消息属于当前会话则直接追加到消息列表，若属于其他联系人则刷新联系人列表。该处理兼顾当前聊天窗口和会话列表的同步。",
        ],
        "功能测试": [
            "测试用例按照核心业务链路组织，重点覆盖登录认证、心理测评、预约咨询、在线聊天和后台统计。每类测试均从输入条件、操作步骤、期望结果和实际表现四个方面进行检查。对于登录测试，重点验证正确账号登录、错误密码提示、未登录访问业务页跳转和角色首页跳转是否符合预期。",
            "测评测试关注量表加载、题目渲染、草稿保存、结果提交和历史记录刷新。预约测试关注用户创建预约、咨询师确认或拒绝、用户取消预约以及已确认预约进入聊天。聊天测试关注联系人加载、历史消息展示、文本消息发送、WebSocket 接收和连接异常提示。管理员测试关注统计卡片、风险分布、测评监控和管理入口跳转。",
        ],
        "兼容性与异常测试": [
            "兼容性测试主要选择 Chrome 和 Edge 浏览器，并在桌面宽屏与窄屏尺寸下观察页面布局。系统采用侧边导航、网格布局和响应式媒体查询，在窄屏下主要内容区域能够转为单列展示，表单、卡片和聊天区域仍保持可读和可操作。",
            "异常测试重点覆盖接口失败、未授权访问、数据为空和配置异常。接口返回 401 时，前端会清理登录状态并跳转至登录页面；列表数据为空时，页面展示空状态而不是空白区域；量表 JSON 解析失败时，系统通过 try-catch 返回空题目列表并提示当前量表未配置变量。",
        ],
    }
    return extras.get(strip_heading_number(title), [])


def add_main(document: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            clean = strip_heading_number(block.text)
            level = 1 if block.level == 2 else 2 if block.level == 3 else 3
            add_heading(document, block.text, level)
            if clean == "业务流程梳理与关键约束":
                add_picture(document, ASSET_DIR / "business-flow.png", "图2.1  系统业务流程图", 5.9)
            if clean == "系统架构设计":
                add_picture(document, ASSET_DIR / "architecture.png", "图3.1  系统总体架构图", 5.9)
            if clean == "数据库概念结构设计":
                add_picture(document, ASSET_DIR / "er-diagram.png", "图3.2  数据库 E-R 图", 5.9)
                add_picture(document, ASSET_DIR / "database-table.png", "表3.1  核心数据表设计", 5.9)
            if clean == "心理风险评估：模型推理与规则回退的双路径":
                add_picture(document, ASSET_DIR / "assessment-flow.png", "图3.3  心理测评风险分析流程图", 5.9)
            if clean == "登录认证与权限控制":
                add_picture(document, ASSET_DIR / "screenshot-login.png", "图4.1  系统登录与角色入口界面", 5.9)
            if clean == "用户端功能实现":
                add_picture(document, ASSET_DIR / "screenshot-user-portal.png", "图4.2  用户端服务首页界面", 5.9)
            if clean == "心理测评与风险分析功能实现":
                add_picture(document, ASSET_DIR / "screenshot-assessment.png", "图4.3  心理测评功能界面", 5.9)
            if clean == "管理员端功能实现":
                add_picture(document, ASSET_DIR / "screenshot-admin-dashboard.png", "图4.4  管理员平台总览界面", 5.9)
            if clean == "在线聊天功能实现":
                add_picture(document, ASSET_DIR / "screenshot-chat.png", "图4.5  在线聊天功能界面", 5.9)
                add_picture(document, ASSET_DIR / "chat-sequence.png", "图4.6  在线聊天模块时序图", 5.9)
            if clean == "不足与展望":
                add_picture(document, ASSET_DIR / "workload-table.png", "表6.1  后续增加工作量方向", 5.9)
            for extra in expansion_after_heading(block.text):
                add_para(document, extra)
            continue
        add_para(document, block.text)


def add_references(document: Document, blocks: list[Block]) -> None:
    for block in blocks:
        if block.kind == "heading":
            add_heading(document, block.text, 1)
            continue
        add_para(document, block.text, first_line=False)


def rewrite_all_footers(docx_path: Path) -> None:
    footer_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:p>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:r><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>
    <w:r><w:fldChar w:fldCharType="separate"/></w:r>
    <w:r><w:t>1</w:t></w:r>
    <w:r><w:fldChar w:fldCharType="end"/></w:r>
  </w:p>
</w:ftr>
""".encode("utf-8")
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as handle:
        temp_path = Path(handle.name)
    with zipfile.ZipFile(docx_path, "r") as src, zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if re.fullmatch(r"word/footer\d+\.xml", item.filename):
                data = footer_xml
            dst.writestr(item, data)
    shutil.move(str(temp_path), docx_path)


def main() -> None:
    create_diagrams()
    blocks = parse_markdown(SOURCE_MD)
    title = title_from(blocks)
    sections = split_sections(blocks)

    template_path = TEMPLATE if TEMPLATE.exists() else FALLBACK_TEMPLATE
    document = Document(str(template_path))
    clear_body(document)
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    add_cover(document, title)
    add_commitment(document)
    add_acknowledgement(document)
    add_abstract(document, sections["abstract_zh"], sections["abstract_en"])
    add_toc_placeholder(document)
    add_main(document, sections["main"])
    add_references(document, sections["references"])
    document.save(OUTPUT)
    rewrite_all_footers(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
