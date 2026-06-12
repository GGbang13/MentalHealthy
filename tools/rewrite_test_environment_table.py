from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from PIL import Image, ImageDraw, ImageFont


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_table_width(table, width_twips: int) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in (1700, 3000, 2200, 3000):
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def insert_paragraph_after(anchor, text: str):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    r_pr.append(sz)
    r.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    anchor.addnext(p)
    return p


def main() -> None:
    src = Path("project_current-1_第五章重排修改版.docx")
    out = Path("project_current-1_测试环境表格版.docx")
    doc = Document(src)

    heading_idx = None
    next_heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == "测试环境":
            heading_idx = i
        elif heading_idx is not None and i > heading_idx and para.style.name.startswith("Heading"):
            next_heading_idx = i
            break
    if heading_idx is None or next_heading_idx is None:
        raise RuntimeError("未找到测试环境小节")

    body = doc.element.body
    heading_el = doc.paragraphs[heading_idx]._element
    for para in doc.paragraphs[heading_idx + 1:next_heading_idx]:
        body.remove(para._element)

    intro = insert_paragraph_after(
        heading_el,
        "为便于说明系统测试所依赖的软硬件与运行条件，测试环境按照前端、后端、数据存储、接口通信和测试关注点进行归纳，具体如表6-1所示。",
    )

    headers = ["测试类别", "环境或工具", "版本/配置", "说明"]
    rows = [
        ("前端运行环境", "Windows、Edge、Node.js、npm", "基于 Vite 启动和构建", "用于运行 Vue 3 前端项目，并检查页面交互、路由跳转和响应式布局。"),
        ("前端主要依赖", "Vue 3、Vue Router、Pinia、Element Plus、Axios、ECharts、TypeScript、Sass", "项目依赖配置", "用于完成页面组件、状态管理、接口请求、图表展示和样式组织。"),
        ("后端运行环境", "Spring Boot 应用服务", "本地开发与测试环境", "用于提供登录认证、测评提交、预约处理、聊天消息和后台统计等 REST API。"),
        ("数据存储环境", "MySQL 数据库", "业务数据表", "用于保存用户、咨询师资料、预约记录、测评记录、聊天消息、文章和通知等数据。"),
        ("实时通信环境", "WebSocket 服务", "/ws/chat", "用于建立聊天长连接，并配合 token 完成身份识别和实时消息推送。"),
        ("接口访问环境", "HTTP 代理与统一接口封装", "/api、/ws/chat", "前端通过 /api 访问后端接口，通过 /ws/chat 建立实时通信连接。"),
        ("测试重点", "功能测试、异常测试、权限测试、兼容性测试", "多角色业务流程", "重点检查角色权限、数据提交、状态流转、异常反馈和页面适配情况。"),
    ]

    image_path = Path("tools/test_environment_table.png")
    image_path.parent.mkdir(parents=True, exist_ok=True)
    make_table_image(image_path, headers, rows)

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Inches(6.2))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run("表6-1 测试环境表")
    cap_run.font.name = "Times New Roman"
    cap_run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap_run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap_run.font.size = Pt(12)

    intro.addnext(pic_para._element)
    pic_para._element.addnext(caption._element)

    doc.save(out)
    print(out)


def wrap_text(draw, text, font, width):
    lines = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def make_table_image(path: Path, headers, rows) -> None:
    width = 2200
    col_widths = [300, 560, 420, 920]
    padding = 18
    try:
        cn_font = ImageFont.truetype("C:/Windows/Fonts/simsun.ttc", 34)
        cn_bold = ImageFont.truetype("C:/Windows/Fonts/simhei.ttf", 34)
    except Exception:
        cn_font = ImageFont.load_default()
        cn_bold = ImageFont.load_default()
    dummy = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(dummy)

    all_rows = [headers] + list(rows)
    wrapped = []
    row_heights = []
    for ridx, row in enumerate(all_rows):
        font = cn_bold if ridx == 0 else cn_font
        wrapped_row = []
        max_lines = 1
        for text, col_w in zip(row, col_widths):
            lines = wrap_text(draw, text, font, col_w - padding * 2)
            wrapped_row.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped.append(wrapped_row)
        row_heights.append(max(68, max_lines * 42 + padding * 2))

    height = sum(row_heights) + 2
    img = Image.new("RGB", (width + 2, height), "white")
    draw = ImageDraw.Draw(img)
    y = 1
    for ridx, (row, row_h) in enumerate(zip(wrapped, row_heights)):
        x = 1
        fill = "#E6E6E6" if ridx == 0 else "#FFFFFF"
        for cidx, (lines, col_w) in enumerate(zip(row, col_widths)):
            draw.rectangle([x, y, x + col_w, y + row_h], fill=fill, outline="black", width=2)
            font = cn_bold if ridx == 0 else cn_font
            text_h = len(lines) * 42
            line_y = y + max(padding, (row_h - text_h) // 2)
            for line in lines:
                if ridx == 0:
                    bbox = draw.textbbox((0, 0), line, font=font)
                    line_x = x + (col_w - (bbox[2] - bbox[0])) / 2
                else:
                    line_x = x + padding
                draw.text((line_x, line_y), line, fill="black", font=font)
                line_y += 42
            x += col_w
        y += row_h
    img.save(path)


if __name__ == "__main__":
    main()
